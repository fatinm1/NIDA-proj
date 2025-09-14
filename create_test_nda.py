#!/usr/bin/env python3
"""
Script to create a test NDA document for testing the AI redlining platform
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_test_nda():
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading('NON-DISCLOSURE AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add parties section
    doc.add_heading('1. PARTIES', level=1)
    doc.add_paragraph('This Non-Disclosure Agreement ("Agreement") is entered into on [DATE] between:')
    
    doc.add_paragraph('Disclosing Party: [FIRM_NAME]')
    doc.add_paragraph('Address: [FIRM_ADDRESS]')
    doc.add_paragraph('City, State ZIP: [FIRM_CITY], [FIRM_STATE] [FIRM_ZIP]')
    
    doc.add_paragraph('')
    doc.add_paragraph('Receiving Party: [CLIENT_NAME]')
    doc.add_paragraph('Address: [CLIENT_ADDRESS]')
    doc.add_paragraph('City, State ZIP: [CLIENT_CITY], [CLIENT_STATE] [CLIENT_ZIP]')
    
    # Add purpose section
    doc.add_heading('2. PURPOSE', level=1)
    doc.add_paragraph('The purpose of this Agreement is to protect confidential and proprietary information that may be disclosed by the Disclosing Party to the Receiving Party in connection with potential business opportunities, including but not limited to:')
    
    doc.add_paragraph('• Technical specifications and designs', style='List Bullet')
    doc.add_paragraph('• Business plans and strategies', style='List Bullet')
    doc.add_paragraph('• Financial information and projections', style='List Bullet')
    doc.add_paragraph('• Customer lists and market data', style='List Bullet')
    doc.add_paragraph('• Any other information marked as confidential', style='List Bullet')
    
    # Add confidential information section
    doc.add_heading('3. CONFIDENTIAL INFORMATION', level=1)
    doc.add_paragraph('For purposes of this Agreement, "Confidential Information" means all non-public, proprietary, or confidential information disclosed by the Disclosing Party to the Receiving Party, whether orally, in writing, or in any other form, including but not limited to:')
    
    doc.add_paragraph('a) Technical data, know-how, research, product plans, products, services, customers, customer lists, markets, software, developments, inventions, processes, formulas, technology, designs, drawings, engineering, hardware configuration information, marketing, finances, or other business information;')
    
    doc.add_paragraph('b) Information that would be considered confidential by a reasonable person in the circumstances;')
    
    doc.add_paragraph('c) Information that is marked, designated, or otherwise identified as "confidential" or "proprietary".')
    
    doc.add_paragraph('Confidential Information does not include information that:')
    doc.add_paragraph('i) Is or becomes publicly available through no breach of this Agreement;', style='List Number')
    doc.add_paragraph('ii) Was rightfully known by the Receiving Party prior to disclosure;', style='List Number')
    doc.add_paragraph('iii) Is rightfully received from a third party without breach of any confidentiality obligation;', style='List Number')
    doc.add_paragraph('iv) Is independently developed by the Receiving Party without use of or reference to the Confidential Information.', style='List Number')
    
    # Add obligations section
    doc.add_heading('4. OBLIGATIONS OF RECEIVING PARTY', level=1)
    doc.add_paragraph('The Receiving Party agrees to:')
    doc.add_paragraph('a) Hold and maintain the Confidential Information in strict confidence;', style='List Number')
    doc.add_paragraph('b) Not disclose the Confidential Information to any third parties without the prior written consent of the Disclosing Party;', style='List Number')
    doc.add_paragraph('c) Use the Confidential Information solely for the purpose of evaluating potential business opportunities;', style='List Number')
    doc.add_paragraph('d) Take reasonable precautions to protect the confidentiality of the Confidential Information;', style='List Number')
    doc.add_paragraph('e) Not make any copies of the Confidential Information except as necessary for the permitted use;', style='List Number')
    doc.add_paragraph('f) Return or destroy all Confidential Information upon request of the Disclosing Party.', style='List Number')
    
    # Add term section
    doc.add_heading('5. TERM', level=1)
    doc.add_paragraph('This Agreement shall remain in effect for a period of five (5) years from the date of execution, and the confidentiality obligations shall survive termination of this Agreement for an additional period of three (3) years.')
    
    # Add return of information section
    doc.add_heading('6. RETURN OF INFORMATION', level=1)
    doc.add_paragraph('Upon termination of this Agreement or upon written request by the Disclosing Party, the Receiving Party shall immediately return or destroy all Confidential Information and all copies thereof, and certify in writing that all such information has been returned or destroyed.')
    
    # Add no license section
    doc.add_heading('7. NO LICENSE', level=1)
    doc.add_paragraph('Nothing in this Agreement shall be construed as granting any rights or licenses to the Receiving Party with respect to any Confidential Information or intellectual property of the Disclosing Party.')
    
    # Add remedies section
    doc.add_heading('8. REMEDIES', level=1)
    doc.add_paragraph('The Receiving Party acknowledges that any breach of this Agreement would cause irreparable harm to the Disclosing Party for which monetary damages would be inadequate. Therefore, the Disclosing Party shall be entitled to seek injunctive relief and other equitable remedies in addition to any other remedies available at law or in equity.')
    
    # Add governing law section
    doc.add_heading('9. GOVERNING LAW', level=1)
    doc.add_paragraph('This Agreement shall be governed by and construed in accordance with the laws of the State of New York, without regard to its conflict of law principles.')
    
    # Add entire agreement section
    doc.add_heading('10. ENTIRE AGREEMENT', level=1)
    doc.add_paragraph('This Agreement constitutes the entire agreement between the parties with respect to the subject matter hereof and supersedes all prior and contemporaneous agreements and understandings, whether written or oral.')
    
    # Add signatures section
    doc.add_heading('11. SIGNATURES', level=1)
    doc.add_paragraph('IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.')
    
    doc.add_paragraph('')
    doc.add_paragraph('DISCLOSING PARTY:')
    doc.add_paragraph('[FIRM_NAME]')
    doc.add_paragraph('')
    doc.add_paragraph('By: [SIGNER_NAME]')
    doc.add_paragraph('Title: [SIGNER_TITLE]')
    doc.add_paragraph('Date: _______________')
    
    doc.add_paragraph('')
    doc.add_paragraph('RECEIVING PARTY:')
    doc.add_paragraph('[CLIENT_NAME]')
    doc.add_paragraph('')
    doc.add_paragraph('By: _________________')
    doc.add_paragraph('Title: _______________')
    doc.add_paragraph('Date: _______________')
    
    # Save the document
    doc.save('test_nda_comprehensive.docx')
    print("✅ Test NDA document created: test_nda_comprehensive.docx")

if __name__ == "__main__":
    create_test_nda()
