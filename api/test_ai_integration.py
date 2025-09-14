#!/usr/bin/env python3
"""
Test script for AI redlining integration
Run this to test the OpenAI GPT-4 integration and document processing
"""

import os
import sys
from dotenv import load_dotenv

# Add the app directory to the path
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from app.services.ai_redlining import AIRedliningService, DocumentProcessor
from docx import Document

def create_test_document():
    """Create a test NDA document for testing"""
    doc = Document()
    
    # Add title
    doc.add_heading('Non-Disclosure Agreement', 0)
    
    # Add parties section
    doc.add_heading('1. Parties', level=1)
    doc.add_paragraph('This Agreement is entered into between [FIRM_NAME] ("Disclosing Party") and the undersigned ("Receiving Party").')
    
    # Add confidentiality section
    doc.add_heading('2. Confidentiality', level=1)
    doc.add_paragraph('The Receiving Party agrees to maintain the confidentiality of all information disclosed by the Disclosing Party for a period of 5 years from the date of disclosure.')
    
    # Add liability section
    doc.add_paragraph('The Receiving Party acknowledges that the Disclosing Party makes no warranties regarding the confidential information.')
    
    # Add signature section
    doc.add_heading('3. Signatures', level=1)
    doc.add_paragraph('IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first above written.')
    doc.add_paragraph('[SIGNER_NAME]')
    doc.add_paragraph('[SIGNER_TITLE]')
    doc.add_paragraph('[FIRM_NAME]')
    
    # Save the test document
    test_file_path = 'test_nda.docx'
    doc.save(test_file_path)
    print(f"Created test document: {test_file_path}")
    
    return test_file_path

def test_ai_service():
    """Test the AI redlining service"""
    print("\n=== Testing AI Redlining Service ===")
    
    # Check if OpenAI API key is set
    if not os.getenv('OPENAI_API_KEY'):
        print("❌ OPENAI_API_KEY not set. Please set it in your .env file")
        return False
    
    try:
        # Initialize AI service
        ai_service = AIRedliningService()
        print("✅ AI service initialized successfully")
        
        # Test document text
        test_text = """
        This is a test NDA document. The confidentiality term is 5 years, which exceeds our standard 2-year limit.
        We need to include investors and agents in the parties section.
        The liability clause should be limited to actual damages only.
        """
        
        # Test custom rules
        custom_rules = [
            {
                'instruction': 'Cap confidentiality obligations at 2 years maximum',
                'category': 'term'
            },
            {
                'instruction': 'Include investors, agents, and financing sources in party definitions',
                'category': 'parties'
            },
            {
                'instruction': 'Limit liability to actual damages and exclude consequential damages',
                'category': 'other'
            }
        ]
        
        print("🔄 Analyzing document with AI...")
        result = ai_service.analyze_document(test_text, custom_rules)
        
        if result['success']:
            print("✅ AI analysis completed successfully")
            print(f"📝 Modifications found: {len(result['redlining_instructions'])}")
            print(f"🤖 AI Analysis: {result['ai_analysis'][:200]}...")
            return True
        else:
            print(f"❌ AI analysis failed: {result['error']}")
            return False
            
    except Exception as e:
        print(f"❌ Error testing AI service: {str(e)}")
        return False

def test_document_processor():
    """Test the document processor"""
    print("\n=== Testing Document Processor ===")
    
    try:
        # Create test document
        test_file_path = create_test_document()
        
        # Initialize processor
        ai_service = AIRedliningService()
        processor = DocumentProcessor(ai_service)
        
        # Test firm details
        firm_details = {
            'name': 'Test Law Firm LLC',
            'address': '123 Test Street',
            'city': 'Test City',
            'state': 'TS',
            'zipCode': '12345',
            'signerName': 'John Test',
            'signerTitle': 'Managing Partner',
            'email': 'john@testfirm.com',
            'phone': '(555) 123-4567'
        }
        
        # Test custom rules
        custom_rules = [
            {
                'instruction': 'Cap confidentiality obligations at 2 years maximum',
                'category': 'term'
            }
        ]
        
        print("🔄 Processing test document...")
        result = processor.process_document(test_file_path, custom_rules, firm_details)
        
        if result['success']:
            print("✅ Document processing completed successfully")
            print(f"📁 Output file: {result['output_path']}")
            print(f"🔧 Modifications applied: {result['modifications_applied']}")
            
            # Clean up test files
            if os.path.exists(test_file_path):
                os.remove(test_file_path)
            if os.path.exists(result['output_path']):
                os.remove(result['output_path'])
            
            return True
        else:
            print(f"❌ Document processing failed: {result['error']}")
            return False
            
    except Exception as e:
        print(f"❌ Error testing document processor: {str(e)}")
        return False

def main():
    """Main test function"""
    print("🧪 Testing AI Redlining Integration")
    print("=" * 50)
    
    # Load environment variables
    load_dotenv()
    
    # Test AI service
    ai_success = test_ai_service()
    
    # Test document processor
    processor_success = test_document_processor()
    
    # Summary
    print("\n" + "=" * 50)
    print("📊 Test Results Summary:")
    print(f"AI Service: {'✅ PASS' if ai_success else '❌ FAIL'}")
    print(f"Document Processor: {'✅ PASS' if processor_success else '❌ FAIL'}")
    
    if ai_success and processor_success:
        print("\n🎉 All tests passed! AI redlining integration is working correctly.")
    else:
        print("\n⚠️  Some tests failed. Please check the error messages above.")
    
    return ai_success and processor_success

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
