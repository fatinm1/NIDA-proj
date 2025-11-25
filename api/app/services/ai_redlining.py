import os
import json
import logging
import uuid
from typing import List, Dict, Any, Optional
from openai import OpenAI
from docx import Document as DocxDocument
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import datetime

logger = logging.getLogger(__name__)

class AIRedliningService:
    def __init__(self):
        try:
            api_key = os.getenv('OPENAI_API_KEY')
            logger.info(f"OpenAI API key status: {'Set' if api_key else 'Not set'}")
            logger.info(f"API key length: {len(api_key) if api_key else 0}")
            logger.info(f"API key starts with: {api_key[:10] if api_key and len(api_key) > 10 else 'N/A'}")
            
            # Check if we should use real API
            if not api_key:
                logger.warning("No OpenAI API key found in environment variables")
                self.client = None
                self.model = "mock-gpt-4"
                logger.warning("Running in mock mode - no OpenAI API calls will be made")
            elif api_key == 'mock-key-for-development':
                logger.warning("Mock API key detected")
                self.client = None
                self.model = "mock-gpt-4"
                logger.warning("Running in mock mode - no OpenAI API calls will be made")
            else:
                # Try to initialize real OpenAI client
                logger.warning("Attempting to initialize OpenAI client with real API key")
                logger.warning(f"API key length: {len(api_key)}")
                logger.warning(f"API key starts with: {api_key[:10]}")
                
                try:
                    # Initialize OpenAI with minimal parameters
                    logger.warning("Creating OpenAI client...")
                    self.client = OpenAI(api_key=api_key)
                    self.model = "gpt-3.5-turbo"
                    logger.warning("OpenAI client initialized successfully with real API")
                        
                except Exception as init_error:
                    logger.error(f"Error during OpenAI client initialization: {str(init_error)}")
                    logger.error(f"Error type: {type(init_error).__name__}")
                    logger.error(f"Error details: {repr(init_error)}")
                    import traceback
                    logger.error(f"Traceback: {traceback.format_exc()}")
                    
                    # Fallback to mock mode
                    self.client = None
                    self.model = "mock-gpt-4"
                    logger.warning("Falling back to mock mode due to OpenAI initialization errors")
        except Exception as e:
            logger.error(f"Error initializing OpenAI client: {str(e)}")
            logger.error(f"Error type: {type(e).__name__}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            self.client = None
            self.model = "mock-gpt-4"
            logger.warning("Falling back to mock mode due to error")
        
    def analyze_document(self, document_text: str, custom_rules: List[Dict[str, Any]], firm_details: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Analyze the document using OpenAI GPT-4 and return redlining instructions
        """
        # Normalize firm_details keys (frontend sends different key names)
        if firm_details:
            normalized = {}
            # Map frontend keys to backend keys
            key_mapping = {
                'name': 'firm_name',
                'signerName': 'signatory_name',
                'signerTitle': 'title',
                'address': 'address',
                'city': 'city',
                'state': 'state',
                'zipCode': 'zip_code',
                'email': 'email',
                'phone': 'phone'
            }
            for frontend_key, backend_key in key_mapping.items():
                if frontend_key in firm_details:
                    normalized[backend_key] = firm_details[frontend_key]
            firm_details = normalized
        
        # Debug: Log the firm details received by AI service
        logger.warning(f"AI Service received firm_details: {firm_details}")
        logger.warning(f"Firm details keys: {list(firm_details.keys()) if firm_details else 'None'}")
        if firm_details:
            for key, value in firm_details.items():
                logger.warning(f"  {key}: '{value}'")
        try:
            # Check if we're in mock mode
            logger.warning(f"AI Service - Client available: {self.client is not None}")
            logger.warning(f"AI Service - Model: {self.model}")
            if not self.client:
                logger.warning("Using mock analysis - OpenAI client not available")
                return self._mock_analysis(document_text, custom_rules, firm_details)
            else:
                logger.warning("Using REAL OpenAI API - client is available")
                logger.warning(f"Client type: {type(self.client)}")
                logger.warning(f"Client attributes: {dir(self.client)}")
            
            # Prepare the prompt for GPT-4
            logger.info("Using REAL OpenAI API for analysis")
            system_prompt = self._build_system_prompt(custom_rules, firm_details)
            user_prompt = self._build_user_prompt(document_text, custom_rules, firm_details)
            
            logger.warning("Making OpenAI API call...")
            try:
                response = self.client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    temperature=0.1,  # Low temperature for consistent legal work
                    max_tokens=2000  # Increased for more detailed responses
                )
                logger.warning("OpenAI API call successful")
            except Exception as api_error:
                logger.error(f"OpenAI API call failed: {str(api_error)}")
                logger.error(f"API error type: {type(api_error).__name__}")
                import traceback
                logger.error(f"API error traceback: {traceback.format_exc()}")
                # Fall back to mock analysis
                logger.warning("Falling back to mock analysis due to API error")
                return self._mock_analysis(document_text, custom_rules, firm_details)
            
            # Parse the AI response
            ai_response = response.choices[0].message.content
            logger.info(f"AI Response: {ai_response}")
            logger.info(f"AI Response length: {len(ai_response)} characters")
            modifications = self._parse_ai_response(ai_response)
            logger.info(f"Parsed modifications: {len(modifications)}")
            
            # Post-process: Ensure firm details are used correctly
            logger.warning(f"POST-PROCESSING: Checking {len(modifications)} modifications for hardcoded values")
            logger.warning(f"POST-PROCESSING: Firm details provided: {firm_details}")
            logger.warning(f"POST-PROCESSING: All modifications from AI:")
            for i, mod in enumerate(modifications):
                logger.warning(f"  Mod {i+1}: {mod.get('type')} - '{mod.get('current_text', 'N/A')[:50]}...' -> '{mod.get('new_text', 'N/A')[:50]}...'")
            
            # VALIDATION: Remove any modifications that incorrectly replace "Company" or dates in wrong contexts
            invalid_modifications = []
            for i, mod in enumerate(modifications):
                current_text = mod.get('current_text', '')
                new_text_val = mod.get('new_text', '')
                
                # Check for date replacements in wrong locations (e.g., document title)
                # Dates should ONLY be in dedicated date fields, nowhere else!
                if any(month in new_text_val for month in ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October', 'November', 'December']):
                    # This is a date replacement - STRICT validation
                    
                    # ONLY ALLOW if current_text is:
                    # 1. Just a date pattern (e.g., "September __, 2025" with nothing else)
                    # 2. A date field label (e.g., "Date: ___")
                    
                    # Check 1: Is it JUST a date pattern with minimal text?
                    current_lower = current_text.lower()
                    forbidden_words = ['company', 'agreement', 'confidentiality', 'nda', 'business', 'dear', 'name', 'for:', 'by:', 'title:', 'effective', 'mutual']
                    is_pure_date = (
                        len(current_text.strip()) < 25 and  # Very short text
                        current_text.count(',') == 1 and  # Has one comma (date format)
                        not any(word in current_lower for word in forbidden_words)
                    )
                    
                    # Check 2: Does it have an explicit date label?
                    has_date_label = current_text.strip().startswith(('Date:', 'Dated:', 'DATE:', 'DATED:'))
                    
                    # REJECT if it doesn't meet either criteria
                    if not is_pure_date and not has_date_label:
                        logger.warning(f"⚠️  REJECTING DATE - Not in a date field!")
                        logger.warning(f"    Text: '{current_text[:100]}'")
                        logger.warning(f"    Length: {len(current_text)}, Pure date: {is_pure_date}, Has label: {has_date_label}")
                        invalid_modifications.append(i)
                        continue
                
                # Check if this modification is trying to replace "Company" in an invalid context
                if 'Company' in current_text:
                    # These are INVALID contexts where Company should NEVER be replaced
                    invalid_contexts = [
                        '(the "Company")',
                        '(the \'Company\')',
                        'the Company',
                        'concerning the Company',
                        'regarding the Company',
                        'about the Company',
                        'of the Company',
                        'business (the',
                        'machining business'
                    ]
                    is_invalid = any(invalid_ctx in current_text for invalid_ctx in invalid_contexts)
                    # VALID contexts are signature blocks only
                    valid_contexts = ['For: Company', 'For:\tCompany', 'For: \tCompany', 'Company (name to be provided upon execution)']
                    is_valid = any(current_text.strip() == valid_ctx or current_text.strip().startswith(valid_ctx) for valid_ctx in valid_contexts)
                    
                    if is_invalid or (not is_valid and len(current_text) > 20):  # Longer text = likely body text, not signature
                        logger.warning(f"⚠️  REJECTING INVALID 'Company' MODIFICATION: '{current_text[:100]}'")
                        invalid_modifications.append(i)
            
            # Remove invalid modifications
            if invalid_modifications:
                logger.warning(f"Removing {len(invalid_modifications)} invalid Company replacements")
                for idx in reversed(invalid_modifications):
                    removed_mod = modifications.pop(idx)
                    logger.warning(f"Removed: {removed_mod}")
            
            if firm_details:
                # Fix any modifications that use hardcoded values instead of firm details
                hardcoded_names = ['John Bagge', 'Jane Doe']
                hardcoded_companies = ['JMC Investment LLC', 'Welch Capital Partners']
                hardcoded_titles = ['Vice President', 'President', 'CEO']
                
                for i, mod in enumerate(modifications):
                    logger.warning(f"POST-PROCESSING Mod {i+1}: {mod}")
                    
                    # FIX: Expand signature fields to include underscores if AI didn't include them
                    current = mod.get('current_text', '').strip()
                    
                    # Check if this is a By:/Title:/Date: modification without underscores
                    if mod.get('type') == 'TEXT_REPLACE':
                        import re
                        
                        # By: field expansion
                        if current == 'By:' or (current.startswith('By:') and len(current) < 10):
                            by_pattern = r'By:[\t\s]+_+'
                            by_matches = list(re.finditer(by_pattern, document_text))
                            if by_matches:
                                by_full_text = by_matches[0].group(0)
                                logger.warning(f"  🔧 EXPANDING 'By:' → '{by_full_text}' (added underscores)")
                                mod['current_text'] = by_full_text
                        
                        # Title: field expansion
                        elif current == 'Title:' or (current.startswith('Title:') and len(current) < 15 and '_' not in current):
                            title_pattern = r'Title:[\t\s]+_+'
                            title_matches = list(re.finditer(title_pattern, document_text))
                            if title_matches:
                                title_full_text = title_matches[0].group(0)
                                logger.warning(f"  🔧 EXPANDING 'Title:' → '{title_full_text}' (added underscores)")
                                mod['current_text'] = title_full_text
                        
                        # Date: field expansion (if needed)
                        elif current == 'Date:' or (current.startswith('Date:') and len(current) < 15 and '_' not in current):
                            date_pattern = r'Date:[\t\s]+_+'
                            date_matches = list(re.finditer(date_pattern, document_text))
                            if date_matches:
                                date_full_text = date_matches[0].group(0)
                                logger.warning(f"  🔧 EXPANDING 'Date:' → '{date_full_text}' (added underscores)")
                                mod['current_text'] = date_full_text
                    
                    # Replace hardcoded names with actual signer name
                    signer_name = firm_details.get('signatory_name') or firm_details.get('signerName')
                    if signer_name:
                        for hardcoded in hardcoded_names:
                            if 'new_text' in mod and hardcoded in mod['new_text']:
                                logger.warning(f"Fixing hardcoded name in modification: '{mod['new_text']}'")
                                mod['new_text'] = mod['new_text'].replace(hardcoded, signer_name)
                                logger.warning(f"Fixed to: '{mod['new_text']}'")
                    
                    # Replace hardcoded companies with actual firm name  
                    company_name = firm_details.get('firm_name') or firm_details.get('name')
                    if company_name:
                        for hardcoded in hardcoded_companies:
                            if 'new_text' in mod and hardcoded in mod['new_text']:
                                logger.warning(f"Fixing hardcoded company in modification: '{mod['new_text']}'")
                                mod['new_text'] = mod['new_text'].replace(hardcoded, company_name)
                                logger.warning(f"Fixed to: '{mod['new_text']}'")
                    
                    # Replace hardcoded titles with actual title
                    title = firm_details.get('title') or firm_details.get('signerTitle')
                    if title:
                        for hardcoded in hardcoded_titles:
                            if 'new_text' in mod and hardcoded in mod['new_text']:
                                logger.warning(f"Fixing hardcoded title in modification: '{mod['new_text']}'")
                                mod['new_text'] = mod['new_text'].replace(hardcoded, title)
                                logger.warning(f"Fixed to: '{mod['new_text']}'")
                
                # Ensure "Dear NAME:" is replaced if it exists
                if firm_details.get('signatory_name') and 'Dear NAME:' in document_text:
                    has_dear_modification = any(
                        mod.get('current_text', '').strip() == 'Dear NAME:' 
                        for mod in modifications
                    )
                    if not has_dear_modification:
                        logger.warning(f"AI didn't generate 'Dear NAME:' modification - adding it manually")
                        modifications.insert(0, {
                            "type": "TEXT_REPLACE",
                            "section": "recipient",
                            "current_text": "Dear NAME:",
                            "new_text": f"Dear {firm_details['signatory_name']}:",
                            "reason": "Replace recipient name placeholder with signer name",
                            "location_hint": "Salutation"
                        })
                
                # Ensure "By:" field is filled - find FULL text with underscores
                if firm_details.get('signatory_name'):
                    import re
                    # Match "By:" followed by tabs/spaces and underscores
                    by_pattern = r'By:[\t\s]+_+'
                    by_matches = list(re.finditer(by_pattern, document_text))
                    if by_matches:
                        # Get the FULL matched text including ALL underscores (don't rstrip!)
                        by_full_text = by_matches[0].group(0)
                        has_by_modification = any(
                            by_full_text in mod.get('current_text', '') or
                            ('By:' in mod.get('current_text', '') and '_' in mod.get('current_text', ''))
                            for mod in modifications
                        )
                        if not has_by_modification:
                            logger.warning(f"Auto-fix: Replacing full By line including underscores")
                            logger.warning(f"  Current: '{by_full_text}'")
                            logger.warning(f"  New: 'By: {firm_details['signatory_name']}'")
                            modifications.append({
                                "type": "TEXT_REPLACE",
                                "section": "signature_block",
                                "current_text": by_full_text,
                                "new_text": f"By: {firm_details['signatory_name']}",
                                "reason": "Fill in signature block with signer name",
                                "location_hint": "Signature block"
                            })
                
                # Ensure "Title:" field is filled - find FULL text with underscores
                if firm_details.get('title'):
                    import re
                    # Match "Title:" followed by tabs/spaces and underscores
                    title_pattern = r'Title:[\t\s]+_+'
                    title_matches = list(re.finditer(title_pattern, document_text))
                    if title_matches:
                        # Get the FULL matched text including ALL underscores (don't rstrip!)
                        title_full_text = title_matches[0].group(0)
                        has_title_modification = any(
                            title_full_text in mod.get('current_text', '') or
                            ('Title:' in mod.get('current_text', '') and '_' in mod.get('current_text', ''))
                            for mod in modifications
                        )
                        if not has_title_modification:
                            logger.warning(f"Auto-fix: Replacing full Title line including underscores")
                            logger.warning(f"  Current: '{title_full_text}'")
                            logger.warning(f"  New: 'Title: {firm_details['title']}'")
                            modifications.append({
                                "type": "TEXT_REPLACE",
                                "section": "signature_block",
                                "current_text": title_full_text,
                                "new_text": f"Title: {firm_details['title']}",
                                "reason": "Fill in signature block with title",
                                "location_hint": "Signature block"
                            })
                
                # Ensure "For: Company" field is filled if it exists
                if firm_details.get('firm_name') and ('For: Company' in document_text or 'For:\tCompany' in document_text or 'For: \tCompany' in document_text):
                    has_for_modification = any(
                        'For:' in mod.get('current_text', '') and 'Company' in mod.get('current_text', '') and firm_details['firm_name'] in mod.get('new_text', '')
                        for mod in modifications
                    )
                    if not has_for_modification:
                        logger.warning(f"AI didn't generate 'For: Company' modification - adding it manually")
                        modifications.append({
                            "type": "TEXT_REPLACE",
                            "section": "signature_block",
                            "current_text": "For: Company",
                            "new_text": f"For: {firm_details['firm_name']}",
                            "reason": "Fill in signature block with firm name",
                            "location_hint": "Signature block"
                        })
            
            # Extract target duration from custom rules
            import re
            target_years = 2  # Default fallback
            term_rule_instruction = ""
            
            for rule in custom_rules:
                rule_name = rule.get('name', '').lower()
                if 'duration' in rule_name or 'term' in rule_name or 'confidentiality' in rule_name:
                    term_rule_instruction = rule.get('instruction', '')
                    year_match = re.search(r'(\d+)\s*(?:\(\d+\))?\s*years?', term_rule_instruction, re.IGNORECASE)
                    if year_match:
                        target_years = int(year_match.group(1))
                        logger.info(f"📅 Extracted target duration from custom rule: {target_years} years")
                        break
            
            target_years_text = f"{target_years} ({target_years}) years"
            target_years_simple = f"{target_years} years"
            
            # Check if document already has the target duration - if so, skip all modifications
            target_patterns_in_doc = [
                target_years_text.lower(),
                target_years_simple.lower(),
                f"{target_years} ({target_years}) year".lower(),  # singular
                f"{target_years} year".lower()  # singular
            ]
            
            # Also check word forms if target is 1-10
            word_to_num = {
                1: 'one', 2: 'two', 3: 'three', 4: 'four', 5: 'five',
                6: 'six', 7: 'seven', 8: 'eight', 9: 'nine', 10: 'ten'
            }
            if target_years in word_to_num:
                word = word_to_num[target_years]
                target_patterns_in_doc.extend([
                    f"{word} ({target_years}) years".lower(),
                    f"{word} years".lower()
                ])
            
            # Check if any target pattern already exists in document
            doc_already_has_target = any(
                pattern in document_text.lower() 
                for pattern in target_patterns_in_doc
            )
            
            if doc_already_has_target:
                logger.info(f"✅ Document already contains target duration '{target_years_text}' - skipping all term modifications")
                # Remove any existing term modifications that would change to the target (already correct)
                modifications = [
                    mod for mod in modifications 
                    if not (mod.get('section') == 'term' and 
                           (target_years_text.lower() in mod.get('new_text', '').lower() or
                            target_years_simple.lower() in mod.get('new_text', '').lower()))
                ]
            else:
                # Ensure "three years" is changed to target duration if it exists
                if 'three years' in document_text.lower():
                    has_term_modification = any(
                        'three years' in mod.get('current_text', '').lower() or 
                        'three (3) years' in mod.get('current_text', '').lower()
                        for mod in modifications
                    )
                    if not has_term_modification:
                        logger.warning(f"AI didn't generate 'three years' modification - adding it manually with target: {target_years} years")
                        # Try to find the exact text in the document
                        if 'three years' in document_text:
                            modifications.append({
                                "type": "TEXT_REPLACE",
                                "section": "term",
                                "current_text": "three years",
                                "new_text": target_years_text,
                                "reason": term_rule_instruction or f"Change confidentiality term to {target_years} years",
                                "location_hint": "Section 13 - Term"
                            })
                        elif 'three (3) years' in document_text:
                            modifications.append({
                                "type": "TEXT_REPLACE",
                                "section": "term",
                                "current_text": "three (3) years",
                                "new_text": target_years_text,
                                "reason": term_rule_instruction or f"Change confidentiality term to {target_years} years",
                                "location_hint": "Section 13 - Term"
                            })
            
            # Auto-fix date placeholders with today's date - be more specific to avoid over-redlining
            import re
            from datetime import datetime
            
            # Common date patterns - be more specific to avoid over-redlining
            # More flexible date pattern recognition - matches various date placeholder formats
            today = datetime.now()
            today_formatted = today.strftime("%B %d, %Y")  # "November 04, 2025" (current date)
            
            # Auto-fix header date specifically (check for any month pattern)
            header_date_pattern = r'([A-Z][a-z]+ _{2,}, \d{4})'
            header_date_matches = re.findall(header_date_pattern, document_text)
            if header_date_matches:
                for header_date_text in header_date_matches:
                    has_header_date_modification = any(
                        header_date_text in mod.get('current_text', '')
                        for mod in modifications
                    )
                    if not has_header_date_modification:
                        logger.warning(f"AI didn't generate header date modification for '{header_date_text}' - adding with today's date")
                        modifications.insert(0, {
                            "type": "TEXT_REPLACE",
                            "section": "header",
                            "current_text": header_date_text,
                            "new_text": today_formatted,
                            "reason": "Fill in header date with today's date",
                            "location_hint": "Document header"
                        })
            
            # Flexible date patterns - using regex to match variations
            flexible_date_patterns = [
                # Month __, Year patterns (any number of underscores)
                (r'[A-Za-z]+ _{2,}, \d{4}', today_formatted),  # "October __, 2025" or "October ___, 2025"
                # Date: _____ patterns
                (r'Date:\s*_+', f"Date: {today_formatted}"),
                # Dated: _____ patterns
                (r'Dated:\s*_+', f"Dated: {today_formatted}"),
                # [DATE] or [date] patterns
                (r'\[DATE\]', today_formatted),
                (r'\[date\]', today_formatted),
                (r'\(DATE\)', today_formatted),
                (r'\(date\)', today_formatted),
                # [Insert date] or [Enter date] patterns
                (r'\[Insert date\]', today_formatted),
                (r'\[Enter date\]', today_formatted),
                (r'\[insert date\]', today_formatted),
                # Month blank, Year - flexible blank matching
                (r'[A-Za-z]+\s+_{2,}\s*,\s*\d{4}', today_formatted),  # Handles spacing variations
            ]
            
            # Search for date patterns in document
            for pattern, replacement_text in flexible_date_patterns:
                matches = re.finditer(pattern, document_text, re.IGNORECASE | re.MULTILINE)
                for match_obj in matches:
                    match_text = match_obj.group(0)
                    # Check if AI already handled this date pattern
                    has_date_modification = any(
                        match_text.lower() in mod.get('current_text', '').lower() or
                        re.sub(r'\s+', ' ', match_text).lower() in re.sub(r'\s+', ' ', mod.get('current_text', '')).lower()
                        for mod in modifications
                    )
                    if not has_date_modification:
                        logger.warning(f"AI didn't generate date modification for pattern '{match_text}' - adding auto-fix: '{match_text}' -> '{replacement_text}'")
                        modifications.append({
                            "type": "TEXT_REPLACE",
                            "section": "date",
                            "current_text": match_text,
                            "new_text": replacement_text,
                            "reason": "Insert today's date using flexible pattern recognition",
                            "location_hint": "Date field"
                        })
            
            return {
                'success': True,
                'redlining_instructions': {
                    'modifications': modifications,
                    'summary': f"AI analysis generated {len(modifications)} modifications",
                    'risk_assessment': "AI-generated risk assessment"
                },
                'ai_analysis': ai_response
            }
            
        except Exception as e:
            logger.error(f"Error in AI analysis: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _mock_analysis(self, document_text: str, custom_rules: List[Dict[str, Any]], firm_details: Dict[str, Any] = None) -> Dict[str, Any]:
        """Provide mock analysis for development/testing"""
        logger.info("Running mock AI analysis")
        logger.info(f"Document text length: {len(document_text)} characters")
        logger.info(f"Number of custom rules: {len(custom_rules)}")
        
        # Create realistic mock redlining instructions based on custom rules
        mock_modifications = []
        
        # Apply firm details first if provided
        if firm_details:
            logger.warning(f"Applying firm details: {firm_details}")
            logger.warning(f"Firm details keys: {list(firm_details.keys())}")
            if 'title' in firm_details:
                logger.warning(f"Title from firm details: '{firm_details['title']}'")
            else:
                logger.warning("No 'title' key in firm details")
        else:
            logger.warning("No firm details provided to mock analysis")
        
        # Check if title field exists in document
        if 'Title:' in document_text:
            logger.warning("Found 'Title:' in document text")
            if 'Title: \t_______________________________' in document_text:
                logger.warning("Found 'Title: \t_______________________________' pattern in document")
            else:
                logger.warning("'Title: \t_______________________________' pattern not found in document")
        else:
            logger.warning("'Title:' not found in document text")
        
        # Look for company name patterns in the document
        company_patterns = [
            "Company (name to be provided upon execution)",
            "For: Company (name to be provided upon execution)",
            "For: Company",
            "Company"
        ]
        
        if firm_details and 'firm_name' in firm_details:
            for pattern in company_patterns:
                if pattern in document_text:
                    if pattern.startswith("For:"):
                        new_text = pattern.replace("Company", firm_details['firm_name'])
                    else:
                        new_text = f"For: {firm_details['firm_name']}"
                    
                    mock_modifications.append({
                        "type": "TEXT_REPLACE",
                        "section": "parties",
                        "current_text": pattern,
                        "new_text": new_text,
                        "reason": "Replace company placeholder with actual firm name",
                        "location_hint": "Parties section"
                    })
                    logger.info(f"Found company pattern: '{pattern}' -> '{new_text}'")
                    break
            else:
                # If no patterns found, add a generic company replacement
                logger.info("No company patterns found, adding generic company replacement")
                mock_modifications.append({
                    "type": "TEXT_REPLACE",
                    "section": "parties",
                    "current_text": "Company",
                    "new_text": f"For: {firm_details['firm_name']}",
                    "reason": "Replace company placeholder with actual firm name",
                    "location_hint": "Parties section"
                })
            
            # Look for signature patterns
            if 'signatory_name' in firm_details:
                if 'By:' in document_text and firm_details['signatory_name'] not in document_text:
                    mock_modifications.append({
                        "type": "TEXT_REPLACE",
                        "section": "signatures",
                        "current_text": "By:",
                        "new_text": f"By: {firm_details['signatory_name']}",
                        "reason": "Replace signer placeholder with actual name",
                        "location_hint": "Signature block"
                    })
                    logger.info(f"Added signer replacement: 'By:' -> 'By: {firm_details['signatory_name']}'")
            
            if 'title' in firm_details:
                logger.warning(f"Processing title field: '{firm_details['title']}'")
                # Look for Title: with various formatting patterns
                title_patterns = [
                    "Title: \t_______________________________",
                    "Title:\t_______________________________",
                    "Title:_______________________________",
                    "Title:",
                    "Title: \t",
                    "Title:\t"
                ]
                
                for title_pattern in title_patterns:
                    logger.warning(f"Checking title pattern: '{title_pattern}'")
                    if title_pattern in document_text:
                        logger.warning(f"Found title pattern '{title_pattern}' in document")
                        if firm_details['title'] not in document_text:
                            logger.warning(f"Title '{firm_details['title']}' not in document, adding replacement")
                            mock_modifications.append({
                                "type": "TEXT_REPLACE",
                                "section": "signatures",
                                "current_text": title_pattern,
                                "new_text": f"Title: {firm_details['title']}",
                                "reason": "Replace title placeholder with actual title",
                                "location_hint": "Signature block"
                            })
                            logger.warning(f"Added title replacement: '{title_pattern}' -> 'Title: {firm_details['title']}'")
                            break
                        else:
                            logger.warning(f"Title '{firm_details['title']}' already in document, skipping replacement")
                    else:
                        logger.warning(f"Title pattern '{title_pattern}' not found in document")
            else:
                logger.warning("No 'title' key in firm_details, skipping title replacement")
        
        for rule in custom_rules:
            rule_name = rule.get('name', '').lower()
            rule_instruction = rule.get('instruction', '')
            logger.info(f"Processing rule: {rule_name}")
            
            # Create specific modifications based on rule type
            if 'duration' in rule_name or 'term' in rule_name or 'confidentiality' in rule_name:
                # Extract target duration from rule instruction
                import re
                target_years = 2  # Default fallback
                
                # Try to extract number from instruction (e.g., "Change to 5 years" → 5)
                year_match = re.search(r'(\d+)\s*(?:\(\d+\))?\s*years?', rule_instruction, re.IGNORECASE)
                if year_match:
                    target_years = int(year_match.group(1))
                    logger.info(f"📅 Extracted target duration: {target_years} years from instruction")
                else:
                    logger.warning(f"⚠️ Could not extract duration from instruction '{rule_instruction}', using default: {target_years} years")
                
                # Build dynamic target pattern
                target_years_text = f"{target_years} ({target_years}) years" if target_years > 0 else f"{target_years} years"
                target_years_simple = f"{target_years} years"
                
                # Check if document already has the target duration - if so, skip modification
                target_patterns_in_doc = [
                    target_years_text.lower(),
                    target_years_simple.lower(),
                    f"{target_years} ({target_years}) year".lower(),  # singular
                    f"{target_years} year".lower()  # singular
                ]
                
                # Also check word forms if target is 1-10
                word_to_num = {
                    1: 'one', 2: 'two', 3: 'three', 4: 'four', 5: 'five',
                    6: 'six', 7: 'seven', 8: 'eight', 9: 'nine', 10: 'ten'
                }
                if target_years in word_to_num:
                    word = word_to_num[target_years]
                    target_patterns_in_doc.extend([
                        f"{word} ({target_years}) years".lower(),
                        f"{word} years".lower()
                    ])
                
                # Check if any target pattern already exists in document
                doc_already_has_target = any(
                    pattern in document_text.lower() 
                    for pattern in target_patterns_in_doc
                )
                
                if doc_already_has_target:
                    logger.info(f"✅ Document already contains target duration '{target_years_text}' - skipping modification")
                    continue  # Skip this rule, document already has the target
                
                # Look for various year patterns in the document and replace with target
                year_patterns = [
                    ('five (5) years', target_years_text),
                    ('5 years', target_years_simple),
                    ('three (3) years', target_years_text),
                    ('3 years', target_years_simple),
                    ('four (4) years', target_years_text),
                    ('4 years', target_years_simple),
                    ('ten (10) years', target_years_text),
                    ('10 years', target_years_simple),
                    ('seven (7) years', target_years_text),
                    ('7 years', target_years_simple),
                    ('three years', target_years_text),
                    ('five years', target_years_text),
                    ('seven years', target_years_text),
                    ('ten years', target_years_text)
                ]
                
                found_pattern = False
                for current_pattern, new_pattern in year_patterns:
                    if current_pattern in document_text.lower():
                        # Double-check: don't replace if it's already the target
                        if current_pattern.lower() == target_years_text.lower() or current_pattern.lower() == target_years_simple.lower():
                            logger.info(f"⚠️ Pattern '{current_pattern}' already matches target '{target_years_text}' - skipping")
                            found_pattern = True  # Mark as found but don't add modification
                            break
                        
                        mock_modifications.append({
                            "type": "TEXT_REPLACE",
                            "section": "term",
                            "current_text": current_pattern,
                            "new_text": new_pattern,
                            "reason": rule_instruction,
                            "location_hint": "Confidentiality term section"
                        })
                        logger.info(f"Found year pattern: {current_pattern} -> {new_pattern}")
                        found_pattern = True
                        break
                
                if not found_pattern:
                    # If no specific patterns found, add a generic year modification
                    logger.info("No specific year patterns found, adding generic modification")
                    mock_modifications.append({
                        "type": "TEXT_REPLACE",
                        "section": "term",
                        "current_text": "three years",
                        "new_text": target_years_text,
                        "reason": rule_instruction,
                        "location_hint": "Confidentiality term section"
                    })
            
            elif 'liability' in rule_name or 'damage' in rule_name:
                # Add liability cap clause
                mock_modifications.append({
                    "type": "TEXT_INSERT",
                    "section": "liability",
                    "new_text": "In no event shall either party be liable for any indirect, incidental, special, consequential, or punitive damages arising out of or relating to this Agreement.",
                    "reason": rule_instruction,
                    "location_hint": "After Section 8, Remedies"
                })
            
            elif 'firm' in rule_name or 'party' in rule_name or 'parties' in rule_name or 'name' in rule_name:
                # Get firm details or use defaults
                firm_name = firm_details.get('firm_name', 'Sample Company LLC') if firm_details else 'Sample Company LLC'
                signer_name = firm_details.get('signatory_name', 'Sample Signer') if firm_details else 'Sample Signer'
                signer_title = firm_details.get('title', 'Authorized Signatory') if firm_details else 'Authorized Signatory'
                
                # Look for party name patterns using firm details
                party_patterns = [
                    ('Company (name to be provided upon execution)', firm_name),
                    ('Recipient', f'{firm_name} (Recipient)'),
                    ('Receiving Party', f'{firm_name} (Recipient)')
                ]
                
                for current_pattern, new_pattern in party_patterns:
                    if current_pattern in document_text:
                        mock_modifications.append({
                            "type": "TEXT_REPLACE",
                            "section": "parties",
                            "current_text": current_pattern,
                            "new_text": new_pattern,
                            "reason": rule_instruction,
                            "location_hint": "Parties section"
                        })
                        logger.info(f"Found party pattern: {current_pattern} -> {new_pattern}")
                        break
                
                # Replace firm placeholders with actual firm details
                if '[FIRM_NAME]' in document_text:
                    mock_modifications.append({
                        "type": "TEXT_REPLACE",
                        "section": "parties",
                        "current_text": "[FIRM_NAME]",
                        "new_text": firm_name,
                        "reason": rule_instruction,
                        "location_hint": "Section 1, line 4"
                    })
                if '[SIGNER_NAME]' in document_text:
                    mock_modifications.append({
                        "type": "TEXT_REPLACE",
                        "section": "signatures",
                        "current_text": "[SIGNER_NAME]",
                        "new_text": signer_name,
                        "reason": rule_instruction,
                        "location_hint": "Section 11, line 55"
                    })
                if '[SIGNER_TITLE]' in document_text:
                    mock_modifications.append({
                        "type": "TEXT_REPLACE",
                        "section": "signatures",
                        "current_text": "[SIGNER_TITLE]",
                        "new_text": signer_title,
                        "reason": rule_instruction,
                        "location_hint": "Section 11, line 56"
                    })
            
            # Add more flexible pattern matching for other rule types
            elif 'governing' in rule_name or 'law' in rule_name:
                # Look for governing law patterns
                law_patterns = [
                    ('State of Delaware', 'State of New York'),
                    ('Delaware', 'New York'),
                    ('Delaware courts', 'New York courts')
                ]
                
                for current_pattern, new_pattern in law_patterns:
                    if current_pattern in document_text:
                        mock_modifications.append({
                            "type": "TEXT_REPLACE",
                            "section": "governing_law",
                            "current_text": current_pattern,
                            "new_text": new_pattern,
                            "reason": rule_instruction,
                            "location_hint": "Governing law section"
                        })
                        logger.info(f"Found law pattern: {current_pattern} -> {new_pattern}")
                        break
            
            elif 'signature' in rule_name or 'block' in rule_name:
                # Get firm details or use defaults
                firm_name = firm_details.get('firm_name', 'Sample Company LLC') if firm_details else 'Sample Company LLC'
                signer_name = firm_details.get('signatory_name', 'Sample Signer') if firm_details else 'Sample Signer'
                signer_title = firm_details.get('title', 'Authorized Signatory') if firm_details else 'Authorized Signatory'
                
                # Replace signature placeholders instead of inserting new content
                signature_replacements = []
                
                # Look for specific signature patterns in the document
                # Try multiple variations of the company placeholder
                company_patterns = [
                    "For: Company (name to be provided upon execution)",
                    "For: Company",
                    "Company (name to be provided upon execution)",
                    "Company"
                ]
                
                for pattern in company_patterns:
                    if pattern in document_text:
                        # Determine the replacement text based on the pattern using firm details
                        if pattern.startswith("For:"):
                            new_text = pattern.replace("Company", firm_name)
                        else:
                            new_text = f"For: {firm_name}"
                        
                        signature_replacements.append({
                            "type": "TEXT_REPLACE",
                            "section": "signatures",
                            "current_text": pattern,
                            "new_text": new_text,
                            "reason": rule_instruction,
                            "location_hint": "Signature block company name"
                        })
                        logger.info(f"Found company pattern: '{pattern}' -> '{new_text}'")
                        break
                
                # Check if signer name already exists to avoid duplicates
                if 'By:' in document_text and signer_name not in document_text:
                    signature_replacements.append({
                        "type": "TEXT_REPLACE",
                        "section": "signatures",
                        "current_text": "By:",
                        "new_text": f"By: {signer_name}",
                        "reason": rule_instruction,
                        "location_hint": "Signature block signer name"
                    })
                
                # Look for Title: with various formatting patterns
                title_patterns = [
                    "Title: \t_______________________________",
                    "Title:\t_______________________________",
                    "Title:_______________________________",
                    "Title:",
                    "Title: \t",
                    "Title:\t"
                ]
                
                for title_pattern in title_patterns:
                    logger.warning(f"Checking rule title pattern: '{title_pattern}'")
                    if title_pattern in document_text:
                        logger.warning(f"Found rule title pattern '{title_pattern}' in document")
                        # Check if title already exists to avoid duplicates
                        if signer_title not in document_text:
                            logger.warning(f"'{signer_title}' not in document, adding replacement")
                            signature_replacements.append({
                                "type": "TEXT_REPLACE",
                                "section": "signatures",
                                "current_text": title_pattern,
                                "new_text": f"Title: {signer_title}",
                                "reason": rule_instruction,
                                "location_hint": "Signature block title"
                            })
                            logger.warning(f"Added rule title replacement: '{title_pattern}' -> 'Title: {signer_title}'")
                            break
                        else:
                            logger.warning(f"'{signer_title}' already in document, skipping replacement")
                    else:
                        logger.warning(f"Rule title pattern '{title_pattern}' not found in document")
                
                if signature_replacements:
                    mock_modifications.extend(signature_replacements)
                    logger.info(f"Added {len(signature_replacements)} signature block replacements")
                else:
                    logger.info("No signature placeholders found to replace")
        
        # If no modifications were found, create a generic one based on the first rule
        if not mock_modifications and custom_rules:
            first_rule = custom_rules[0]
            mock_modifications.append({
                "type": "TEXT_INSERT",
                "section": "general",
                "new_text": f"[Modification based on rule: {first_rule.get('name', 'Unknown')}]",
                "reason": first_rule.get('instruction', 'Rule application'),
                "location_hint": "Document beginning"
            })
            logger.info("Added fallback modification")
        
        logger.info(f"Mock analysis complete. Generated {len(mock_modifications)} modifications:")
        for i, mod in enumerate(mock_modifications):
            logger.info(f"  {i+1}. {mod['type']}: '{mod['current_text']}' -> '{mod['new_text']}'")
        
        return {
            'success': True,
            'redlining_instructions': {
                'modifications': mock_modifications,
                'summary': f"Mock analysis completed with {len(mock_modifications)} modifications based on {len(custom_rules)} rules",
                'risk_assessment': "Mock assessment - modifications based on provided rules for testing purposes"
            },
            'ai_analysis': f"Mock AI analysis generated {len(mock_modifications)} redlining suggestions based on the provided rules. In production, this would be generated by OpenAI GPT-4."
        }
    
    def _build_system_prompt(self, custom_rules: List[Dict[str, Any]], firm_details: Dict[str, Any] = None) -> str:
        """Build the system prompt for GPT-4"""
        base_prompt = """You are an expert legal AI assistant specializing in NDA (Non-Disclosure Agreement) redlining. 
        Your task is to analyze NDA documents and provide PRECISE, TARGETED redlining instructions.
        
        CRITICAL INSTRUCTIONS:
        1. Make ONLY the specific changes requested in the rules - do NOT over-redline
        2. Use EXACT text matches from the document for current_text
        3. Focus on the specific areas mentioned in the rules
        4. **FIRM DETAILS OVERRIDE RULES**: When firm details are provided, ALWAYS use them instead of any company names or person names mentioned in the rules
        5. Be EXTREMELY conservative - only change what is explicitly requested
        6. REPLACE existing placeholders - do NOT create new fields or sections
        7. For signature blocks, ONLY use TEXT_REPLACE - NEVER use TEXT_INSERT
        8. Replace "By:" with "By: [Firm Details Signer Name]", "Title:" with "Title: [Firm Details Title]", "For:" with "For: [Firm Details Company Name]"
        9. Do NOT add new signature lines or duplicate existing ones
        10. If a rule says "JMC Investment LLC" but firm details provide "MyCompany", USE "MyCompany"
        11. **CRITICAL - NEVER TOUCH THESE**: Do NOT replace "Company" when it appears in legal text:
            - (the "Company") - this is a legal definition
            - "the Company" - refers to the disclosing party
            - "concerning the Company" 
            - "the Company acknowledges"
            - "acquisition of our client's machining business (the 'Company')"
            - ANY instance where "Company" is in quotes or parentheses in the body text
        12. **ONLY REPLACE "Company" IN SIGNATURE BLOCKS**: 
            - "For: Company" → "For: [FIRM_NAME]"
            - "For:\tCompany" → "For:\t[FIRM_NAME]"
            - "Company (name to be provided upon execution)" → "[FIRM_NAME]"
        13. **ABSOLUTELY NEVER** replace "Company" in the first few paragraphs where it defines what "Company" means
        14. If you're unsure whether to replace "Company", DON'T replace it - only replace in signature blocks at the end
        15. When changing years/durations, ONLY make changes in numbered sections about Term/Duration, NOT in the opening paragraphs
        16. When a term duration rule is provided, extract the target duration from the rule instruction (e.g., "Change to 5 years" means replace with "5 (5) years")
        17. **FLEXIBLE DATE PATTERN RECOGNITION**: For date insertion rules, use pattern recognition, NOT exact text matching
        18. Recognize ANY date placeholder pattern (blanks, underscores, brackets, "[DATE]", etc.) and replace with today's date
        19. Be flexible with date patterns - "Month __, Year", "Month ___, Year", "[DATE]", "Date: _______" all should be recognized and replaced
        
        REDLINING PRINCIPLES:
        - Replace specific text with exact matches
        - Insert new text only where specified
        - Delete text only when explicitly requested
        - Maintain legal document structure
        - Use professional legal language
        
        Available modification types:
        - TEXT_REPLACE: Replace specific text (most common)
        - TEXT_INSERT: Insert new text at specific locations
        - TEXT_DELETE: Remove specific text
        - CLAUSE_ADD: Add entire new clauses
        
        Return your analysis in this exact JSON format (use the ACTUAL firm details values provided above, NOT these example values):
        {
            "modifications": [
                {
                    "type": "TEXT_REPLACE",
                    "section": "confidentiality_term",
                    "current_text": "three years",
                    "new_text": "X (X) years",
                    "reason": "Change term duration as specified in custom rule",
                    "location_hint": "Term section"
                }
            ],
            "summary": "Brief summary of all changes",
            "risk_assessment": "Assessment of any legal risks in modifications"
        }
        
        IMPORTANT: For signature blocks, always use TEXT_REPLACE to fill in existing placeholders. Never use TEXT_INSERT to create new signature fields."""
        
        if custom_rules:
            rules_text = "\n\nCUSTOM RULES TO APPLY (follow these exactly):\n"
            for rule in custom_rules:
                instruction = rule['instruction']
                
                # CRITICAL: Replace any placeholders or hardcoded values in rules with actual firm details
                original_instruction = instruction
                if firm_details:
                    # Replace placeholder tokens with actual firm details
                    if '[FIRM_NAME]' in instruction and firm_details.get('firm_name'):
                        instruction = instruction.replace('[FIRM_NAME]', firm_details['firm_name'])
                    
                    if '[SIGNER_NAME]' in instruction and firm_details.get('signatory_name'):
                        instruction = instruction.replace('[SIGNER_NAME]', firm_details['signatory_name'])
                    
                    if '[TITLE]' in instruction and firm_details.get('title'):
                        instruction = instruction.replace('[TITLE]', firm_details['title'])
                    
                    # Also replace any remaining hardcoded company names
                    hardcoded_companies = ['JMC Investment LLC', 'JMC Investment', 'JMC', 'Welch Capital Partners']
                    for company in hardcoded_companies:
                        if company in instruction and firm_details.get('firm_name'):
                            instruction = instruction.replace(company, firm_details['firm_name'])
                    
                    # Replace hardcoded person names
                    hardcoded_names = ['John Bagge', 'John', 'Jane Doe']
                    for name in hardcoded_names:
                        if name in instruction and firm_details.get('signatory_name'):
                            instruction = instruction.replace(name, firm_details['signatory_name'])
                    
                    # Replace hardcoded titles
                    hardcoded_titles = ['Vice President', 'President', 'CEO', 'Managing Director']
                    for title in hardcoded_titles:
                        if title in instruction and firm_details.get('title'):
                            instruction = instruction.replace(title, firm_details['title'])
                    
                    # Log if instruction was modified
                    if instruction != original_instruction:
                        logger.warning(f"RULE TRANSFORMED:")
                        logger.warning(f"  BEFORE: {original_instruction}")
                        logger.warning(f"  AFTER:  {instruction}")
                
                rules_text += f"- {instruction}\n"
            base_prompt += rules_text
        
        if firm_details:
            firm_text = "\n\n" + "="*80 + "\n"
            firm_text += "🚨 CRITICAL - MANDATORY FIRM DETAILS - DO NOT USE ANY OTHER VALUES 🚨\n"
            firm_text += "="*80 + "\n"
            firm_text += "YOU MUST USE THESE EXACT VALUES - DO NOT SUBSTITUTE WITH EXAMPLES:\n\n"
            
            if firm_details.get('firm_name'):
                firm_text += f"COMPANY NAME TO USE: '{firm_details['firm_name']}'\n"
                firm_text += f"⚠️  **LOCATION**: ONLY at the END of the document in the signature block\n"
                firm_text += f"✅ CORRECT: Find 'For: Company' (in signature block at end) → Replace with 'For: {firm_details['firm_name']}'\n"
                firm_text += f"✅ CORRECT: Find 'For:\tCompany' (in signature block at end) → Replace with 'For:\t{firm_details['firm_name']}'\n"
                firm_text += f"❌ WRONG: Do NOT replace (the \"Company\") in the opening paragraphs\n"
                firm_text += f"❌ WRONG: Do NOT replace \"the Company\" anywhere in the legal text body\n"
                firm_text += f"❌ WRONG: Do NOT replace \"Company\" if it's in quotes or parentheses\n"
                firm_text += f"  - In JSON: current_text should be 'For: Company' (the exact text from signature block only)\n\n"
            
            if firm_details.get('signatory_name'):
                firm_text += f"SIGNER NAME TO USE: '{firm_details['signatory_name']}'\n"
                firm_text += f"  - Find 'Dear NAME:' in document → Replace entire text with 'Dear {firm_details['signatory_name']}:'\n"
                firm_text += f"  - Find 'By:' (with blank or underscore) → Replace with 'By: {firm_details['signatory_name']}'\n"
                firm_text += f"  - In JSON: current_text should be the PLACEHOLDER ('Dear NAME:' or 'By:'), not the firm detail value\n"
                firm_text += f"  - DO NOT search for '{firm_details['signatory_name']}' in the document!\n\n"
            
            if firm_details.get('title'):
                firm_text += f"TITLE TO USE: '{firm_details['title']}'\n"
                firm_text += f"  - Find 'Title:' with blank/underscore → Replace with 'Title: {firm_details['title']}'\n"
                firm_text += f"  - In JSON: current_text should be 'Title: \\t_______________________________' (the PLACEHOLDER)\n"
                firm_text += f"  - DO NOT search for '{firm_details['title']}' in the document!\n\n"
            
            firm_text += "="*80 + "\n"
            firm_text += "REMINDER: Use the EXACT values above. Do NOT use placeholder examples.\n"
            firm_text += "="*80 + "\n"
            base_prompt += firm_text
        
        return base_prompt
    
    def _build_user_prompt(self, document_text: str, custom_rules: List[Dict[str, Any]], firm_details: Dict[str, Any] = None) -> str:
        """Build the user prompt with document content"""
        # Use more of the document text for better analysis
        text_limit = 4000  # Increased from 2000
        document_preview = document_text[:text_limit]
        
        prompt = f"""Please analyze the following NDA document and provide PRECISE redlining instructions based on the custom rules.

DOCUMENT CONTENT:
{document_preview}

CRITICAL REQUIREMENTS:
1. When providing current_text, copy the EXACT text as it appears in the document above
2. For example, if the document says "five (5) years", your current_text must be "five (5) years", not "5 years" or "five years"
3. Look carefully at the exact wording, punctuation, and formatting
4. Make ONLY the changes specified in the rules - be conservative
5. Focus on the specific areas mentioned in the custom rules
6. **MOST IMPORTANT**: If firm details are provided below, use them instead of ANY company/person names mentioned in the rules
7. **CRITICAL**: Do NOT replace "Company" when it refers to the disclosing party in legal text
8. **CRITICAL**: For date placeholders (e.g., "Month __, Year"), replace with today's current date in "Month Day, Year" format
9. **CRITICAL**: Only replace placeholders, not actual content

        COMMON PATTERNS TO LOOK FOR:
        - "five (5) years" (most common in legal documents)
        - "five years" 
        - "5 years"
        - "five (5) year"
        - "five year"
        - "three years" or "three (3) years"
        - "Company (name to be provided upon execution)" or just "Company"
        - "For: Company" or "For: Company (name to be provided upon execution)"
        - "Dear NAME:" or "Dear [Name]:" (for recipient names)
        - "State of Delaware" or "Delaware"
        - Signature blocks: "By:", "Title: \t_______________________________", "For: Company"
        - "Title: \t_______________________________" (most common title pattern)
        
        FLEXIBLE DATE PATTERN RECOGNITION:
        The AI should intelligently recognize ANY date placeholder pattern and replace it with today's date.
        This includes but is NOT limited to:
        
        Pattern Examples (use regex-like flexibility):
        - "Month __, Year" or "Month __, YYYY" → "Month Day, Year" (e.g., "October __, 2025" → "October 27, 2025")
        - "Month ___, Year" → "Month Day, Year" (any number of underscores)
        - "[DATE]", "[date]", "(DATE)", "(date)" → "Month Day, Year"
        - "Date: _______________" or "Date: __________" → "Date: Month Day, Year"
        - "Dated: _______________" or "Dated: __________" → "Dated: Month Day, Year"
        - "_______" following "Date" or "Dated" → Replace with today's date
        - "Month __, YYYY" in header/footer → "Month Day, YYYY"
        - Any text like "[Insert date]" or "[Enter date]" → "Month Day, Year"
        - Pattern: "Month [blank], Year" where [blank] can be "__", "___,", "[ ]", "_______", etc.
        
        CRITICAL DATE RULES:
        - If you see ANY pattern that requests a date insertion (blanks, underscores, brackets, "[DATE]", etc.), replace it
        - Do NOT wait for exact pattern matches - use pattern recognition
        - If a date field is incomplete or has placeholders, fill it with today's date
        - ONLY replace dates that are incomplete/placeholder - DO NOT change complete dates
        - Use today's date in "Month Day, Year" format (e.g., "October 27, 2025")
        - Use pattern matching, not exact text matching - be flexible with variations

CRITICAL: When replacing text, use the EXACT text as it appears in the document.

IMPORTANT: Use the FULL PHRASE as current_text, not just the placeholder word.
For example:
- Use "Dear NAME:" not just "NAME"
- Use "For: Company" not just "Company"
- Use "By:" not just "By"
- Use "Title:" not just "Title"

        SIGNATURE BLOCK HANDLING:
        - ONLY use TEXT_REPLACE for signature blocks - NEVER use TEXT_INSERT
        - Replace "By:" with "By: [SIGNER_NAME]" using firm details provided
        - Replace "Title: \t_______________________________" with "Title: [SIGNER_TITLE]" using firm details
        - Replace "Title:\t_______________________________" with "Title: [SIGNER_TITLE]" using firm details
        - Replace "For: Company" with "For: [FIRM_NAME]" using firm details
        - Replace "Company (name to be provided upon execution)" with "[FIRM_NAME]" using firm details
        - Replace "Dear NAME:" with "Dear [SIGNER_NAME]:" using firm details
        - Replace just "Company" with "[FIRM_NAME]" if it's in a "For:" context
        - Do NOT insert new signature blocks or create new fields
        - Do NOT duplicate existing signature lines
        - Find existing placeholders and replace them in place
        - ALWAYS use the actual firm details values provided, NOT these placeholder examples

IMPORTANT: Only make the specific changes requested in the custom rules. Do not over-redline or make unnecessary changes.

Please provide your analysis in the specified JSON format."""

        # Add firm details reminder at the end for maximum emphasis
        if firm_details:
            prompt += f"\n\n" + "!"*80 + "\n"
            prompt += f"🚨 FINAL REMINDER - MANDATORY REPLACEMENTS - USE THESE EXACT VALUES:\n"
            prompt += "!"*80 + "\n"
            if firm_details.get('signatory_name'):
                prompt += f'✅ REQUIRED: Replace "Dear NAME:" with "Dear {firm_details["signatory_name"]}:"\n'
                prompt += f'✅ REQUIRED: Replace "By:" with "By: {firm_details["signatory_name"]}"\n'
            if firm_details.get('firm_name'):
                prompt += f'✅ REQUIRED: Replace "For: Company" with "For: {firm_details["firm_name"]}"\n'
            if firm_details.get('title'):
                prompt += f'✅ REQUIRED: Replace "Title:" fields with "Title: {firm_details["title"]}"\n'
            prompt += f"\n❌ DO NOT use placeholder examples - use actual firm details only!\n"
            prompt += f"❌ DO NOT leave placeholders like 'Dear NAME:', 'Company', or blanks unchanged!\n"
            prompt += "!"*80 + "\n"

        return prompt
    
    def _parse_ai_response(self, ai_response: str) -> List[Dict[str, Any]]:
        """Parse the AI response and extract redlining instructions"""
        try:
            # Try to extract JSON from the response
            if '{' in ai_response and '}' in ai_response:
                start = ai_response.find('{')
                end = ai_response.rfind('}') + 1
                json_str = ai_response[start:end]
                
                # Fix: Escape control characters (tabs, newlines) that might be in the JSON
                # This is necessary because the AI might include literal tabs in the response
                import re
                # Replace literal tab characters with escaped tabs
                json_str = json_str.replace('\t', '\\t')
                # Replace literal newlines with escaped newlines (if any)
                json_str = re.sub(r'(?<!\\)\n', '\\n', json_str)
                
                logger.info(f"Parsing AI response JSON (length: {len(json_str)})")
                parsed = json.loads(json_str)
                modifications = parsed.get('modifications', [])
                
                # Unescape control characters in the modifications
                # After JSON parsing, \t becomes literal "\t" but we need actual tab characters
                for mod in modifications:
                    if 'current_text' in mod:
                        # Replace escaped sequences with actual characters
                        mod['current_text'] = mod['current_text'].replace('\\t', '\t').replace('\\n', '\n')
                    if 'new_text' in mod:
                        mod['new_text'] = mod['new_text'].replace('\\t', '\t').replace('\\n', '\n')
                
                logger.info(f"Successfully parsed {len(modifications)} modifications from AI response")
                return modifications
            else:
                logger.warning("Could not find JSON in AI response")
                return []
                
        except json.JSONDecodeError as e:
            logger.error(f"JSON parsing error: {str(e)}")
            logger.error(f"JSON string that failed to parse: {json_str[:500] if 'json_str' in locals() else 'N/A'}")
            return []
        except Exception as e:
            logger.error(f"Error parsing AI response: {str(e)}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            return []

    def generate_changes_for_review(self, document_text: str, custom_rules: List[Dict], firm_details: Dict[str, str]) -> Dict[str, Any]:
        """Generate changes with unique IDs for review interface"""
        logger.info("Generating changes for review interface")
        
        # Get AI modifications - analyze_document returns a dict, extract the list
        analysis_result = self.analyze_document(document_text, custom_rules, firm_details)
        
        # Extract modifications from the nested structure
        if isinstance(analysis_result, dict):
            if 'redlining_instructions' in analysis_result and isinstance(analysis_result['redlining_instructions'], dict):
                modifications = analysis_result['redlining_instructions'].get('modifications', [])
            elif 'modifications' in analysis_result:
                modifications = analysis_result['modifications']
            else:
                logger.error(f"Unexpected analysis_result structure: {analysis_result.keys()}")
                modifications = []
        elif isinstance(analysis_result, list):
            # If it's already a list, use it directly
            modifications = analysis_result
        else:
            logger.error(f"Unexpected analysis_result type: {type(analysis_result)}")
            modifications = []
        
        logger.warning(f"Extracted {len(modifications)} modifications for review interface")
        
        # Add unique IDs and metadata to each change
        changes = []
        for i, mod in enumerate(modifications):
            # Ensure mod is a dictionary, not a string
            if isinstance(mod, str):
                logger.warning(f"Skipping invalid modification (string): {mod}")
                continue
            if not isinstance(mod, dict):
                logger.warning(f"Skipping invalid modification (not dict): {type(mod)}")
                continue
                
            change_id = str(uuid.uuid4())
            change = {
                "id": change_id,
                "type": mod.get("type", "TEXT_REPLACE"),
                "section": mod.get("section", "general"),
                "current_text": mod.get("current_text", ""),
                "new_text": mod.get("new_text", ""),
                "reason": mod.get("reason", ""),
                "location_hint": mod.get("location_hint", ""),
                "status": "pending",  # pending, accepted, rejected
                "order": i + 1
            }
            changes.append(change)
        
        return {
            "changes": changes,
            "total_changes": len(changes),
            "document_text": document_text,
            "firm_details": firm_details,
            "custom_rules": custom_rules
        }

    def apply_accepted_changes(self, doc_path: str, accepted_changes: List[Dict], signature_path: str = None) -> Dict[str, Any]:
        """Apply only the accepted changes to create final document"""
        accepted_count = len([c for c in accepted_changes if c.get("status") == "accepted"])
        logger.warning(f"========== APPLYING {accepted_count} ACCEPTED CHANGES ==========")
        logger.warning(f"Total changes received: {len(accepted_changes)}")
        
        try:
            # Load the document
            doc = DocxDocument(doc_path)
            logger.warning(f"Document loaded: {doc_path}")
            logger.warning(f"Document has {len(doc.paragraphs)} paragraphs")
            
            # Enable Track Changes mode in the document settings
            try:
                settings = doc.settings
                settings_element = settings.element
                # Add trackRevisions setting if it doesn't exist
                track_changes = settings_element.find(qn('w:trackRevisions'))
                if track_changes is None:
                    track_changes = OxmlElement('w:trackRevisions')
                    settings_element.append(track_changes)
                logger.warning("Track Changes mode enabled in document settings")
            except Exception as tc_error:
                logger.warning(f"Could not enable Track Changes mode: {tc_error}")
            
            # Apply each accepted change
            changes_actually_applied = 0
            for idx, change in enumerate(accepted_changes):
                if change.get("status") == "accepted":
                    logger.warning(f"\nApplying change {idx+1}/{len(accepted_changes)}:")
                    logger.warning(f"  Old: '{change.get('current_text', '')[:80]}'")
                    logger.warning(f"  New: '{change.get('new_text', '')[:80]}'")
                    result = self._apply_single_change(doc, change)
                    if result:
                        changes_actually_applied += 1
                        logger.warning(f"  ✅ Applied successfully")
                    else:
                        logger.warning(f"  ❌ Failed to apply")
            
            logger.warning(f"Successfully applied {changes_actually_applied} out of {accepted_count} accepted changes")
            
            # Apply signature if provided
            if signature_path and os.path.exists(signature_path):
                self._apply_signature(doc, signature_path)
            
            # Generate output path
            output_path = self._generate_output_path(doc_path)
            logger.warning(f"Saving final document to: {output_path}")
            
            # Save the final document
            doc.save(output_path)
            logger.warning(f"Document saved successfully")
            
            return {
                "success": True,
                "output_path": output_path,
                "changes_applied": changes_actually_applied
            }
            
        except Exception as e:
            logger.error(f"Error applying accepted changes: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def _generate_output_path(self, input_path: str) -> str:
        """Generate output path for the processed document"""
        # Create outputs directory if it doesn't exist
        output_dir = 'outputs'
        os.makedirs(output_dir, exist_ok=True)
        
        # Get the absolute path to the outputs directory
        base_dir = os.path.abspath(output_dir)
        filename = os.path.basename(input_path)
        name, ext = os.path.splitext(filename)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{name}_final_{timestamp}{ext}"
        return os.path.abspath(os.path.join(base_dir, output_filename))
    
    def _apply_signature(self, doc: DocxDocument, signature_path: str):
        """Apply signature image to the document"""
        try:
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            logger.info(f"Applying signature from: {signature_path}")
            
            # Find "Signed:" text and add signature right after it
            signature_added = False
            
            for paragraph in doc.paragraphs:
                if 'Signed:' in paragraph.text:
                    logger.info(f"Found 'Signed:' in paragraph: {paragraph.text}")
                    
                    # Clear the paragraph and rebuild it with proper signature placement
                    original_text = paragraph.text
                    paragraph.clear()
                    
                    # Split text around "Signed:"
                    parts = original_text.split('Signed:', 1)
                    
                    # Add text before "Signed:"
                    if parts[0].strip():
                        before_run = paragraph.add_run(parts[0])
                    
                    # Add "Signed:" text
                    signed_run = paragraph.add_run("Signed: ")
                    
                    # Add signature image immediately after "Signed:"
                    signature_run = paragraph.add_run()
                    signature_run.add_picture(signature_path, width=Inches(1.2))
                    
                    # Add any remaining text after "Signed:"
                    if len(parts) > 1 and parts[1].strip():
                        remaining_run = paragraph.add_run(parts[1])
                    
                    signature_added = True
                    logger.info("Signature added right next to 'Signed:' text")
                    break
            
            # If no "Signed:" found, look for signature placeholders
            if not signature_added:
                for paragraph in doc.paragraphs:
                    if '[SIGNATURE]' in paragraph.text:
                        # Replace placeholder with signature
                        paragraph.text = paragraph.text.replace('[SIGNATURE]', '')
                        
                        # Add signature image
                        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                        run.add_picture(signature_path, width=Inches(1.5))
                        signature_added = True
                        logger.info("Signature added at [SIGNATURE] placeholder")
                        break
            
            # If still not added, add at the end of the document
            if not signature_added:
                logger.warning("No signature placeholder found, adding at end of document")
                last_paragraph = doc.paragraphs[-1]
                run = last_paragraph.add_run()
                run.add_picture(signature_path, width=Inches(1.5))
                
        except Exception as e:
            logger.error(f"Error applying signature: {str(e)}")
    
    def _apply_single_change(self, doc: DocxDocument, change: Dict) -> bool:
        """Apply a single change to the document, returns True if successful"""
        try:
            old_text = change.get("current_text", "")
            new_text = change.get("new_text", "")
            
            if not old_text or not new_text:
                logger.warning(f"Skipping change with empty text: {change}")
                return False
            
            # Find and replace text in the document
            found = False
            for para_idx, paragraph in enumerate(doc.paragraphs):
                # Try exact match first
                if old_text in paragraph.text:
                    logger.warning(f"  Found exact match in paragraph {para_idx}: '{paragraph.text[:100]}'")
                    success = self._replace_text_in_paragraph(paragraph, old_text, new_text)
                    if success:
                        logger.warning(f"  ✅ Change applied in paragraph {para_idx}")
                        found = True
                        break
                    else:
                        logger.warning(f"  ❌ Failed to apply change in paragraph {para_idx}")
            
            # If not found, try with whitespace normalization (e.g., "For: Company" vs "For:\tCompany")
            if not found:
                import re
                # Normalize whitespace for comparison
                old_text_normalized = re.sub(r'\s+', ' ', old_text.strip())
                
                for para_idx, paragraph in enumerate(doc.paragraphs):
                    para_text_normalized = re.sub(r'\s+', ' ', paragraph.text.strip())
                    if old_text_normalized in para_text_normalized:
                        # Find the actual text in the original paragraph
                        # Extract the actual text with original whitespace
                        actual_old_text = self._find_text_with_whitespace(paragraph.text, old_text)
                        if actual_old_text:
                            logger.warning(f"  Found with whitespace variation in paragraph {para_idx}")
                            logger.warning(f"    Looking for: '{old_text}'")
                            logger.warning(f"    Found actual: '{actual_old_text}'")
                            success = self._replace_text_in_paragraph(paragraph, actual_old_text, new_text)
                            if success:
                                logger.warning(f"  ✅ Change applied with whitespace normalization in paragraph {para_idx}")
                                found = True
                                break
            
            if not found:
                logger.warning(f"  ⚠️  Text not found in any paragraph: '{old_text[:80]}'")
                logger.warning(f"  Try checking the document for: variations with tabs/spaces")
                return False
            
            return True
                    
        except Exception as e:
            logger.error(f"Error applying single change: {str(e)}")
            return False
    
    def _find_text_with_whitespace(self, paragraph_text: str, search_text: str) -> str:
        """Find text in paragraph that matches after whitespace normalization"""
        import re
        
        # Try exact match first
        if search_text in paragraph_text:
            return search_text
        
        # Try common whitespace variations
        variations = [
            search_text.replace(' ', '\t'),  # Replace spaces with tabs
            search_text.replace(' ', '  '),  # Replace single space with double
            re.sub(r'\s+', '\t', search_text),  # All whitespace to tabs
        ]
        
        for variation in variations:
            if variation in paragraph_text:
                logger.warning(f"    Found with whitespace variation: '{variation}'")
                return variation
        
        # Try pattern matching with flexible whitespace
        # For "For: Company", match "For:<any whitespace>Company"
        pattern = re.escape(search_text)
        pattern = pattern.replace(r'\ ', r'\s+')  # Allow any whitespace where there was a space
        
        match = re.search(pattern, paragraph_text)
        if match:
            logger.warning(f"    Found with pattern matching: '{match.group(0)}'")
            return match.group(0)
        
        return None
    
    def _replace_text_in_paragraph(self, paragraph, old_text: str, new_text: str) -> bool:
        """Replace text in a paragraph while showing BOTH old and new text with Track Changes"""
        try:
            # Always use paragraph-level replacement to show BOTH old and new text
            if old_text in paragraph.text:
                logger.warning(f"    Applying Track Changes to show old and new text")
                
                # Store original text
                original_text = paragraph.text
                
                # Clear all runs and create new ones
                paragraph.clear()
                
                # Split text around the old_text and create runs
                parts = original_text.split(old_text, 1)  # Split only once for first occurrence
                logger.warning(f"    Text split into {len(parts)} parts")
                
                # Add text before the replacement (no formatting)
                if parts[0]:
                    before_run = paragraph.add_run(parts[0])
                    logger.warning(f"    Added text before: '{parts[0][:50]}'")
                
                # Add OLD text with BLACK strikethrough (shows what was removed)
                deleted_run = paragraph.add_run(old_text)
                deleted_run.font.strike = True
                deleted_run.font.color.rgb = RGBColor(0, 0, 0)  # Black strikethrough
                logger.warning(f"    ✅ Added DELETION (black strikethrough): '{old_text[:50]}'")
                
                # Add NEW text with RED underline (shows what was added)
                added_run = paragraph.add_run(new_text)
                added_run.font.underline = True
                added_run.font.color.rgb = RGBColor(255, 0, 0)  # Red underline for new text
                logger.warning(f"    ✅ Added INSERTION (red underline): '{new_text[:50]}'")
                
                # Add any remaining text after the replacement (no formatting)
                if len(parts) > 1 and parts[1]:
                    after_run = paragraph.add_run(parts[1])
                    logger.warning(f"    Added text after: '{parts[1][:50]}'")
                
                logger.warning(f"    ✅✅ Track Changes complete - document will show BOTH old (strikethrough) and new (red underline) text")
                return True
            else:
                logger.warning(f"Text '{old_text}' not found in paragraph: '{paragraph.text}'")
                # Try case-insensitive search with change tracking
                if old_text.lower() in paragraph.text.lower():
                    logger.info(f"Found case-insensitive match, trying replacement with change tracking")
                    
                    # Find the actual old text in the paragraph
                    text_lower = paragraph.text.lower()
                    idx = text_lower.find(old_text.lower())
                    
                    if idx != -1:
                        # Extract the actual old text from the paragraph
                        actual_old_text = paragraph.text[idx:idx+len(old_text)]
                        
                        # Clear the paragraph and rebuild with change tracking
                        original_text = paragraph.text
                        paragraph.clear()
                        
                        # Split on the actual old text
                        parts = original_text.split(actual_old_text, 1)
                        
                        # Add text before replacement
                        if parts[0]:
                            run = paragraph.add_run(parts[0])
                        
                        # Add OLD text with black strikethrough
                        deleted_run = paragraph.add_run(actual_old_text)
                        deleted_run.font.strike = True
                        deleted_run.font.color.rgb = RGBColor(0, 0, 0)  # Black strikethrough
                        
                        # Add NEW text with red underline
                        added_run = paragraph.add_run(new_text)
                        added_run.font.underline = True
                        added_run.font.color.rgb = RGBColor(255, 0, 0)  # Red underline
                        
                        # Add remaining text
                        if len(parts) > 1:
                            run = paragraph.add_run(parts[1])
                        
                        logger.info(f"Case-insensitive replacement with visual change tracking: '{actual_old_text}' (strikethrough) -> '{new_text}' (red underline)")
                        return True
                
        except Exception as e:
            logger.error(f"Error replacing text in paragraph: {str(e)}")
            # Fallback to simple replacement with change tracking
            try:
                original_text = paragraph.text
                if old_text in original_text:
                    paragraph.clear()
                    parts = original_text.split(old_text, 1)
                    
                    if parts[0]:
                        paragraph.add_run(parts[0])
                    
                        # Use native Word Track Changes for deletion
                        deleted_run = paragraph.add_run(old_text)
                        self._add_track_change_deletion(deleted_run)
                        
                        # Use native Word Track Changes for insertion
                        added_run = paragraph.add_run(new_text)
                        self._add_track_change_insertion(added_run)
                    
                    if len(parts) > 1:
                        paragraph.add_run(parts[1])
                    
                    logger.info(f"Fallback replacement with visual change tracking successful: '{old_text}' (strikethrough) -> '{new_text}' (red underline)")
                    return True
                else:
                    logger.warning(f"Text '{old_text}' not found in fallback replacement")
                    return False
            except Exception as fallback_error:
                logger.error(f"Fallback replacement failed: {str(fallback_error)}")
                return False
        
        return False
    
    def _add_track_change_deletion(self, run):
        """Add Word Track Changes deletion markup to a run"""
        try:
            import uuid
            # Create deletion markup - w:del must contain w:r (run) elements
            del_elem = OxmlElement('w:del')
            change_id = str(uuid.uuid4())[:8]  # Use unique ID for each change
            del_elem.set(qn('w:id'), change_id)
            del_elem.set(qn('w:author'), 'AI Redlining System')
            del_elem.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
            
            # Create a run element inside the deletion
            run_elem = OxmlElement('w:r')
            
            # Add run properties (formatting)
            rPr = OxmlElement('w:rPr')
            # Copy any formatting from the original run
            if run.font.bold:
                b_elem = OxmlElement('w:b')
                rPr.append(b_elem)
            if run.font.italic:
                i_elem = OxmlElement('w:i')
                rPr.append(i_elem)
            # Add strikethrough for deleted text
            strike_elem = OxmlElement('w:strike')
            rPr.append(strike_elem)
            run_elem.append(rPr)
            
            # Use w:delText for deleted text (required for Track Changes deletions)
            text_elem = OxmlElement('w:delText')
            text_elem.set(qn('xml:space'), 'preserve')  # Preserve whitespace
            text_elem.text = run.text
            run_elem.append(text_elem)
            del_elem.append(run_elem)
            
            # Replace the run's XML with the deletion markup
            run._element.getparent().replace(run._element, del_elem)
            logger.warning(f"      ✅ Track Changes deletion XML created")
            
        except Exception as e:
            logger.warning(f"      ❌ Track Changes deletion failed: {e}")
            logger.warning(f"      Using fallback strikethrough")
            # Fallback to strikethrough
            run.font.strike = True
            run.font.color.rgb = RGBColor(0, 0, 0)
    
    def _add_track_change_insertion(self, run):
        """Add Word Track Changes insertion markup to a run"""
        try:
            import uuid
            # Create insertion markup - w:ins must contain w:r (run) elements
            ins_elem = OxmlElement('w:ins')
            change_id = str(uuid.uuid4())[:8]  # Use unique ID for each change
            ins_elem.set(qn('w:id'), change_id)
            ins_elem.set(qn('w:author'), 'AI Redlining System')
            ins_elem.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
            
            # Create a run element inside the insertion
            run_elem = OxmlElement('w:r')
            
            # Add run properties (formatting)
            rPr = OxmlElement('w:rPr')
            # Add underline for inserted text
            u_elem = OxmlElement('w:u')
            u_elem.set(qn('w:val'), 'single')
            rPr.append(u_elem)
            # Add color for inserted text (red)
            color_elem = OxmlElement('w:color')
            color_elem.set(qn('w:val'), 'FF0000')  # Red
            rPr.append(color_elem)
            run_elem.append(rPr)
            
            # Add the text to the run element
            text_elem = OxmlElement('w:t')
            text_elem.set(qn('xml:space'), 'preserve')  # Preserve whitespace
            text_elem.text = run.text
            run_elem.append(text_elem)
            ins_elem.append(run_elem)
            
            # Replace the run's XML with the insertion markup
            run._element.getparent().replace(run._element, ins_elem)
            logger.warning(f"      ✅ Track Changes insertion XML created")
            
        except Exception as e:
            logger.warning(f"      ❌ Track Changes insertion failed: {e}")
            logger.warning(f"      Using fallback red underline")
            # Fallback to red underline
            run.font.underline = True
            run.font.color.rgb = RGBColor(255, 0, 0)

class DocumentProcessor:
    def __init__(self, ai_service: AIRedliningService):
        self.ai_service = ai_service
    
    def process_document(self, doc_path: str, custom_rules: List[Dict[str, Any]], 
                        firm_details: Dict[str, Any], signature_path: str = None) -> Dict[str, Any]:
        """
        Process a Word document with AI redlining and return the modified document
        Optimized for large files with memory efficiency
        """
        try:
            # Check file size and use appropriate processing method
            file_size = os.path.getsize(doc_path)
            logger.info(f"Processing document: {file_size} bytes ({file_size/1024:.1f} KB)")
            
            # For large files (>100KB), use chunked processing
            if file_size > 100 * 1024:  # 100KB threshold
                return self._process_large_document(doc_path, custom_rules, firm_details, signature_path)
            else:
                return self._process_small_document(doc_path, custom_rules, firm_details, signature_path)
            
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _process_small_document(self, doc_path: str, custom_rules: List[Dict[str, Any]], 
                               firm_details: Dict[str, Any], signature_path: str = None) -> Dict[str, Any]:
        """Process small documents (<100KB) using standard method"""
        try:
            # Load the document
            doc = DocxDocument(doc_path)
            
            # Extract text for AI analysis
            document_text = self._extract_document_text(doc)
            
            # Get AI redlining instructions
            ai_result = self.ai_service.analyze_document(document_text, custom_rules, firm_details)
            
            if not ai_result['success']:
                return {
                    'success': False,
                    'error': ai_result['error']
                }
            
            # Apply AI modifications
            modifications = ai_result['redlining_instructions']
            if isinstance(modifications, dict) and 'modifications' in modifications:
                modifications = modifications['modifications']
            
            logger.info(f"AI generated {len(modifications) if modifications else 0} modifications")
            if modifications:
                for i, mod in enumerate(modifications):
                    logger.info(f"Modification {i+1}: {mod}")
            else:
                logger.warning("No modifications generated by AI")
                logger.info(f"AI result: {ai_result}")
            
            # Apply modifications if we have any
            if modifications:
                self._apply_modifications(doc, modifications)
            
            # Apply firm details
            self._apply_firm_details(doc, firm_details)
            
            # Apply signature if provided
            if signature_path and os.path.exists(signature_path):
                self._apply_signature(doc, signature_path)
            
            # Generate output path
            output_path = self._generate_output_path(doc_path)
            
            # Ensure output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Save the modified document
            doc.save(output_path)
            
            return {
                'success': True,
                'output_path': output_path,
                'modifications_applied': len(modifications) if modifications else 0,
                'ai_analysis': ai_result.get('ai_analysis', 'No analysis available')
            }
            
        except Exception as e:
            logger.error(f"Error processing small document: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _process_large_document(self, doc_path: str, custom_rules: List[Dict[str, Any]], 
                               firm_details: Dict[str, Any], signature_path: str = None) -> Dict[str, Any]:
        """Process large documents (>100KB) using memory-efficient chunked method"""
        try:
            logger.info("Processing large document with chunked method")
            
            # Load the document
            doc = DocxDocument(doc_path)
            
            # Extract text in chunks for AI analysis
            document_text = self._extract_document_text_chunked(doc)
            
            # Get AI redlining instructions
            ai_result = self.ai_service.analyze_document(document_text, custom_rules, firm_details)
            
            if not ai_result['success']:
                return {
                    'success': False,
                    'error': ai_result['error']
                }
            
            # Apply AI modifications
            modifications = ai_result['redlining_instructions']
            if isinstance(modifications, dict) and 'modifications' in modifications:
                modifications = modifications['modifications']
            
            # Apply modifications if we have any
            if modifications:
                self._apply_modifications_chunked(doc, modifications)
            
            # Apply firm details
            self._apply_firm_details_chunked(doc, firm_details)
            
            # Apply signature if provided
            if signature_path and os.path.exists(signature_path):
                self._apply_signature(doc, signature_path)
            
            # Generate output path
            output_path = self._generate_output_path(doc_path)
            
            # Ensure output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Save the modified document
            doc.save(output_path)
            
            return {
                'success': True,
                'output_path': output_path,
                'modifications_applied': len(modifications) if modifications else 0,
                'ai_analysis': ai_result.get('ai_analysis', 'No analysis available')
            }
            
        except Exception as e:
            logger.error(f"Error processing large document: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _extract_document_text(self, doc: DocxDocument) -> str:
        """Extract text content from the Word document"""
        text_parts = []
        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)
        
        return '\n'.join(text_parts)
    
    def _extract_document_text_chunked(self, doc: DocxDocument) -> str:
        """Extract text content from large Word documents in chunks for memory efficiency"""
        text_parts = []
        chunk_size = 1000  # Process 1000 paragraphs at a time
        total_paragraphs = len(doc.paragraphs)
        
        logger.info(f"Extracting text from {total_paragraphs} paragraphs in chunks of {chunk_size}")
        
        for i in range(0, total_paragraphs, chunk_size):
            chunk_end = min(i + chunk_size, total_paragraphs)
            chunk_text = []
            
            for j in range(i, chunk_end):
                if j < len(doc.paragraphs):
                    chunk_text.append(doc.paragraphs[j].text)
            
            text_parts.extend(chunk_text)
            
            # Log progress for large documents
            if total_paragraphs > 5000:
                progress = (chunk_end / total_paragraphs) * 100
                logger.info(f"Text extraction progress: {progress:.1f}%")
        
        return '\n'.join(text_parts)
    
    def _apply_modifications(self, doc: DocxDocument, modifications: List[Dict[str, Any]]):
        """Apply AI-generated modifications to the document"""
        for mod in modifications:
            try:
                if mod['type'] == 'TEXT_REPLACE':
                    self._replace_text(doc, mod['current_text'], mod['new_text'])
                elif mod['type'] == 'TEXT_INSERT':
                    self._insert_text(doc, mod['new_text'], mod.get('location_hint', ''))
                elif mod['type'] == 'TEXT_DELETE':
                    self._delete_text(doc, mod['current_text'])
                elif mod['type'] == 'CLAUSE_ADD':
                    self._add_clause(doc, mod['new_text'])
                    
            except Exception as e:
                logger.warning(f"Failed to apply modification {mod}: {str(e)}")
                continue
    
    def _apply_modifications_chunked(self, doc: DocxDocument, modifications: List[Dict[str, Any]]):
        """Apply AI-generated modifications to large documents in chunks for memory efficiency"""
        logger.info(f"Applying {len(modifications)} modifications to large document")
        
        for i, mod in enumerate(modifications):
            try:
                if mod['type'] == 'TEXT_REPLACE':
                    self._replace_text_chunked(doc, mod['current_text'], mod['new_text'])
                elif mod['type'] == 'TEXT_INSERT':
                    self._insert_text_chunked(doc, mod['new_text'], mod.get('location_hint', ''))
                elif mod['type'] == 'TEXT_DELETE':
                    self._delete_text_chunked(doc, mod['current_text'])
                elif mod['type'] == 'CLAUSE_ADD':
                    self._add_clause_chunked(doc, mod['new_text'])
                
                # Log progress for large documents
                if len(modifications) > 10:
                    progress = ((i + 1) / len(modifications)) * 100
                    logger.info(f"Modification progress: {progress:.1f}%")
                    
            except Exception as e:
                logger.warning(f"Failed to apply modification {mod}: {str(e)}")
                continue
    
    def _replace_text(self, doc: DocxDocument, old_text: str, new_text: str):
        """Replace text in the document with professional redlining"""
        logger.info(f"Attempting to replace '{old_text}' with '{new_text}'")
        logger.info(f"Document has {len(doc.paragraphs)} paragraphs")
        replaced = False
        
        # Log first few paragraphs to see what text is actually in the document
        logger.warning("First 10 paragraphs in document:")
        for i, para in enumerate(doc.paragraphs[:10]):
            logger.warning(f"Paragraph {i+1}: '{para.text}'")
        
        # Also search for common patterns to see what's actually in the document
        logger.warning("Searching for common patterns in document:")
        common_patterns = ["For:", "Company", "Dear", "NAME", "Title:", "By:"]
        for pattern in common_patterns:
            found_paragraphs = []
            for i, para in enumerate(doc.paragraphs):
                if pattern.lower() in para.text.lower():
                    found_paragraphs.append(f"Para {i+1}: '{para.text}'")
            if found_paragraphs:
                logger.warning(f"Found '{pattern}' in: {found_paragraphs}")
            else:
                logger.warning(f"'{pattern}' not found in document")
        
        # Search for variations of the specific text we're trying to replace
        logger.warning(f"Searching for variations of '{old_text}':")
        variations_to_search = [
            old_text,
            old_text.replace(" ", ""),  # No spaces
            old_text.replace(":", ""),  # No colon
            old_text.lower(),           # Lowercase
            old_text.upper(),           # Uppercase
            old_text.replace("Company", "Comapny"),  # Common typo
            old_text.replace("Company", "company"),  # Case variation
            "For: Company",  # Common pattern
            "For:Company",   # No space after colon
            "For Company",   # No colon
            "for: company",  # All lowercase
            "FOR: COMPANY",  # All uppercase
            # Handle multiple whitespace variations
            "For:  Company",  # Double space
            "For:   Company", # Triple space
            "For:    Company", # Quadruple space
            "For:\tCompany",  # Tab after colon
            "For: \tCompany", # Space and tab
            "For:\t Company", # Tab and space
            "For: \t Company", # Space, tab, space
        ]
        for variation in variations_to_search:
            found_paragraphs = []
            for i, para in enumerate(doc.paragraphs):
                if variation.lower() in para.text.lower():
                    found_paragraphs.append(f"Para {i+1}: '{para.text}'")
            if found_paragraphs:
                logger.warning(f"Found variation '{variation}' in: {found_paragraphs}")
            else:
                logger.warning(f"Variation '{variation}' not found in document")
        
        # First, try exact match
        for paragraph in doc.paragraphs:
            if old_text in paragraph.text:
                logger.warning(f"Found exact match in paragraph: {paragraph.text}")
                if self._replace_text_in_paragraph(paragraph, old_text, new_text):
                    logger.warning(f"Successfully replaced. New paragraph: {paragraph.text}")
                    replaced = True
                    break
        
        # If exact match failed, try normalized whitespace matching
        if not replaced:
            logger.warning("Trying normalized whitespace matching")
            import re
            # Normalize whitespace in the old text (replace multiple spaces/tabs with single space)
            normalized_old = re.sub(r'\s+', ' ', old_text.strip())
            for paragraph in doc.paragraphs:
                # Normalize whitespace in paragraph text
                normalized_para = re.sub(r'\s+', ' ', paragraph.text.strip())
                if normalized_old.lower() in normalized_para.lower():
                    logger.warning(f"Found normalized match in paragraph: {paragraph.text}")
                    # Try to find the actual text in the paragraph and replace it
                    if old_text in paragraph.text:
                        if self._replace_text_in_paragraph(paragraph, old_text, new_text):
                            logger.warning(f"Successfully replaced with normalized matching. New paragraph: {paragraph.text}")
                            replaced = True
                            break
                    else:
                        # If exact text not found, try to find and replace the normalized version
                        # Find the actual text pattern in the paragraph
                        for variation in variations_to_search:
                            if variation in paragraph.text:
                                logger.warning(f"Found variation '{variation}' in paragraph, replacing with '{new_text}'")
                                if self._replace_text_in_paragraph(paragraph, variation, new_text):
                                    logger.warning(f"Successfully replaced variation. New paragraph: {paragraph.text}")
                                    replaced = True
                                    break
                        if replaced:
                            break
        
        if not replaced:
            logger.warning(f"Exact text '{old_text}' not found, trying variations")
            # Try partial matches for common variations
            variations = [
                old_text.replace("5", "five (5)"),
                old_text.replace("five (5)", "5"),
                old_text.replace("years", "year"),
                old_text.replace("year", "years"),
                "five (5) years",  # Common legal format
                "five years",
                "5 year",
                "five year",
                # For company name variations
                old_text.replace("For: ", ""),
                old_text.replace("Company (name to be provided upon execution)", "Company"),
                "Company",
                "For: Company",
                "Company (name to be provided upon execution)",
                "For: Company (name to be provided upon execution)",
                # Additional variations for better matching
                "For: " + old_text,
                old_text + " (name to be provided upon execution)",
                "For: " + old_text + " (name to be provided upon execution)"
            ]
            
            for variation in variations:
                if variation and variation != old_text:  # Skip empty or duplicate variations
                    logger.info(f"Trying variation: '{variation}'")
                    for paragraph in doc.paragraphs:
                        if variation in paragraph.text:
                            logger.info(f"Found variation '{variation}' in paragraph: {paragraph.text}")
                            if self._replace_text_in_paragraph(paragraph, variation, new_text):
                                logger.info(f"Successfully replaced variation. New paragraph: {paragraph.text}")
                                replaced = True
                                break
                    if replaced:
                        break
            
            # If still not found, try case-insensitive search
            if not replaced:
                logger.info("Trying case-insensitive search")
                for paragraph in doc.paragraphs:
                    if old_text.lower() in paragraph.text.lower():
                        logger.info(f"Found case-insensitive match in paragraph: {paragraph.text}")
                        if self._replace_text_in_paragraph(paragraph, old_text, new_text):
                            logger.info(f"Successfully replaced case-insensitive. New paragraph: {paragraph.text}")
                            replaced = True
                            break
        
        if not replaced:
            logger.error(f"Failed to replace '{old_text}' with '{new_text}' - no matches found")
        else:
            logger.info(f"Successfully replaced '{old_text}' with '{new_text}'")
    
    def _replace_text_in_paragraph(self, paragraph, old_text: str, new_text: str):
        """Replace text in a paragraph while preserving formatting"""
        try:
            # First, try to find and replace in individual runs
            for run in paragraph.runs:
                if old_text in run.text:
                    # Replace the text in the run
                    run.text = run.text.replace(old_text, new_text)
                    
                    # Add redlining formatting
                    run.font.underline = True
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red for additions
                    logger.info(f"Replaced text in run: '{old_text}' -> '{new_text}'")
                    return True
            
            # If not found in individual runs, try paragraph-level replacement
            if old_text in paragraph.text:
                logger.info(f"Text found in paragraph, doing paragraph-level replacement")
                
                # Store original text
                original_text = paragraph.text
                
                # Clear all runs and create new ones
                paragraph.clear()
                
                # Split text around the old_text and create runs
                parts = original_text.split(old_text)
                logger.info(f"Split into {len(parts)} parts: {parts}")
                
                # Add text before the replacement
                if parts[0]:  # Add non-empty parts
                    run = paragraph.add_run(parts[0])
                    # Keep original formatting (no change tracking)
                
                # Use native Word Track Changes for deletion
                deleted_run = paragraph.add_run(old_text)
                self._add_track_change_deletion(deleted_run)
                
                # Use native Word Track Changes for insertion
                added_run = paragraph.add_run(new_text)
                self._add_track_change_insertion(added_run)
                
                # Add any remaining text after the replacement
                if len(parts) > 1:
                    run = paragraph.add_run(parts[1])
                    # Keep original formatting
                
                logger.info(f"Visual change tracking added: '{old_text}' (strikethrough) -> '{new_text}' (red underline)")
                
                logger.info(f"Final paragraph text: {paragraph.text}")
                return True
            else:
                logger.warning(f"Text '{old_text}' not found in paragraph: '{paragraph.text}'")
                # Try case-insensitive search with change tracking
                if old_text.lower() in paragraph.text.lower():
                    logger.info(f"Found case-insensitive match, trying replacement with change tracking")
                    
                    # Find the actual old text in the paragraph
                    text_lower = paragraph.text.lower()
                    idx = text_lower.find(old_text.lower())
                    
                    if idx != -1:
                        # Extract the actual old text from the paragraph
                        actual_old_text = paragraph.text[idx:idx+len(old_text)]
                        
                        # Clear the paragraph and rebuild with change tracking
                        original_text = paragraph.text
                        paragraph.clear()
                        
                        # Split on the actual old text
                        parts = original_text.split(actual_old_text, 1)
                        
                        # Add text before replacement
                        if parts[0]:
                            run = paragraph.add_run(parts[0])
                        
                        # Add OLD text with black strikethrough
                        deleted_run = paragraph.add_run(actual_old_text)
                        deleted_run.font.strike = True
                        deleted_run.font.color.rgb = RGBColor(0, 0, 0)  # Black strikethrough
                        
                        # Add NEW text with red underline
                        added_run = paragraph.add_run(new_text)
                        added_run.font.underline = True
                        added_run.font.color.rgb = RGBColor(255, 0, 0)  # Red underline
                        
                        # Add remaining text
                        if len(parts) > 1:
                            run = paragraph.add_run(parts[1])
                        
                        logger.info(f"Case-insensitive replacement with visual change tracking: '{actual_old_text}' (strikethrough) -> '{new_text}' (red underline)")
                        return True
                
        except Exception as e:
            logger.error(f"Error replacing text in paragraph: {str(e)}")
            # Fallback to simple replacement with change tracking
            try:
                original_text = paragraph.text
                if old_text in original_text:
                    paragraph.clear()
                    parts = original_text.split(old_text, 1)
                    
                    if parts[0]:
                        paragraph.add_run(parts[0])
                    
                        # Use native Word Track Changes for deletion
                        deleted_run = paragraph.add_run(old_text)
                        self._add_track_change_deletion(deleted_run)
                        
                        # Use native Word Track Changes for insertion
                        added_run = paragraph.add_run(new_text)
                        self._add_track_change_insertion(added_run)
                    
                    if len(parts) > 1:
                        paragraph.add_run(parts[1])
                    
                    logger.info(f"Fallback replacement with visual change tracking successful: '{old_text}' (strikethrough) -> '{new_text}' (red underline)")
                    return True
                else:
                    logger.warning(f"Text '{old_text}' not found in fallback replacement")
                    return False
            except Exception as fallback_error:
                logger.error(f"Fallback replacement failed: {str(fallback_error)}")
                return False
        
        return False
    
    def _add_track_change_deletion(self, run):
        """Add Word Track Changes deletion markup to a run"""
        try:
            import uuid
            # Create deletion markup - w:del must contain w:r (run) elements
            del_elem = OxmlElement('w:del')
            change_id = str(uuid.uuid4())[:8]  # Use unique ID for each change
            del_elem.set(qn('w:id'), change_id)
            del_elem.set(qn('w:author'), 'AI Redlining System')
            del_elem.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
            
            # Create a run element inside the deletion
            run_elem = OxmlElement('w:r')
            
            # Add run properties (formatting)
            rPr = OxmlElement('w:rPr')
            # Copy any formatting from the original run
            if run.font.bold:
                b_elem = OxmlElement('w:b')
                rPr.append(b_elem)
            if run.font.italic:
                i_elem = OxmlElement('w:i')
                rPr.append(i_elem)
            # Add strikethrough for deleted text
            strike_elem = OxmlElement('w:strike')
            rPr.append(strike_elem)
            run_elem.append(rPr)
            
            # Use w:delText for deleted text (required for Track Changes deletions)
            text_elem = OxmlElement('w:delText')
            text_elem.set(qn('xml:space'), 'preserve')  # Preserve whitespace
            text_elem.text = run.text
            run_elem.append(text_elem)
            del_elem.append(run_elem)
            
            # Replace the run's XML with the deletion markup
            run._element.getparent().replace(run._element, del_elem)
            logger.warning(f"      ✅ Track Changes deletion XML created")
            
        except Exception as e:
            logger.warning(f"      ❌ Track Changes deletion failed: {e}")
            logger.warning(f"      Using fallback strikethrough")
            # Fallback to strikethrough
            run.font.strike = True
            run.font.color.rgb = RGBColor(0, 0, 0)
    
    def _add_track_change_insertion(self, run):
        """Add Word Track Changes insertion markup to a run"""
        try:
            import uuid
            # Create insertion markup - w:ins must contain w:r (run) elements
            ins_elem = OxmlElement('w:ins')
            change_id = str(uuid.uuid4())[:8]  # Use unique ID for each change
            ins_elem.set(qn('w:id'), change_id)
            ins_elem.set(qn('w:author'), 'AI Redlining System')
            ins_elem.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
            
            # Create a run element inside the insertion
            run_elem = OxmlElement('w:r')
            
            # Add run properties (formatting)
            rPr = OxmlElement('w:rPr')
            # Add underline for inserted text
            u_elem = OxmlElement('w:u')
            u_elem.set(qn('w:val'), 'single')
            rPr.append(u_elem)
            # Add color for inserted text (red)
            color_elem = OxmlElement('w:color')
            color_elem.set(qn('w:val'), 'FF0000')  # Red
            rPr.append(color_elem)
            run_elem.append(rPr)
            
            # Add the text to the run element
            text_elem = OxmlElement('w:t')
            text_elem.set(qn('xml:space'), 'preserve')  # Preserve whitespace
            text_elem.text = run.text
            run_elem.append(text_elem)
            ins_elem.append(run_elem)
            
            # Replace the run's XML with the insertion markup
            run._element.getparent().replace(run._element, ins_elem)
            logger.warning(f"      ✅ Track Changes insertion XML created")
            
        except Exception as e:
            logger.warning(f"      ❌ Track Changes insertion failed: {e}")
            logger.warning(f"      Using fallback red underline")
            # Fallback to red underline
            run.font.underline = True
            run.font.color.rgb = RGBColor(255, 0, 0)
    
    def _replace_text_chunked(self, doc: DocxDocument, old_text: str, new_text: str):
        """Replace text in large documents with chunked processing for memory efficiency"""
        logger.info(f"Attempting to replace '{old_text}' with '{new_text}' in large document")
        replaced = False
        chunk_size = 500  # Process 500 paragraphs at a time
        total_paragraphs = len(doc.paragraphs)
        
        for i in range(0, total_paragraphs, chunk_size):
            chunk_end = min(i + chunk_size, total_paragraphs)
            
            for j in range(i, chunk_end):
                if j < len(doc.paragraphs):
                    paragraph = doc.paragraphs[j]
                    if old_text in paragraph.text:
                        logger.info(f"Found text to replace in paragraph {j}: {paragraph.text[:100]}...")
                        
                        # Replace text while preserving formatting
                        self._replace_text_in_paragraph(paragraph, old_text, new_text)
                        replaced = True
                        break
            
            if replaced:
                break
            
            # Log progress for very large documents
            if total_paragraphs > 10000:
                progress = (chunk_end / total_paragraphs) * 100
                logger.info(f"Text replacement progress: {progress:.1f}%")
        
        if not replaced:
            logger.warning(f"Text '{old_text}' not found in large document for replacement")
    
    def _insert_text(self, doc: DocxDocument, text: str, location_hint: str):
        """Insert new text at specified location"""
        # Find appropriate location and insert
        new_paragraph = doc.add_paragraph(text)
        for run in new_paragraph.runs:
            # Professional redlining for insertions
            run.font.underline = True
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red for additions
    
    def _insert_text_chunked(self, doc: DocxDocument, text: str, location_hint: str):
        """Insert new text at specified location in large documents"""
        # Find appropriate location and insert
        new_paragraph = doc.add_paragraph(text)
        for run in new_paragraph.runs:
            # Professional redlining for insertions
            run.font.underline = True
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red for additions
    
    def _delete_text(self, doc: DocxDocument, text: str):
        """Delete text from the document"""
        for paragraph in doc.paragraphs:
            if text in paragraph.text:
                paragraph.text = paragraph.text.replace(text, '')
                # Add strikethrough formatting for deleted text
                for run in paragraph.runs:
                    run.font.strike = True
    
    def _delete_text_chunked(self, doc: DocxDocument, text: str):
        """Delete text from large documents with chunked processing"""
        chunk_size = 500
        total_paragraphs = len(doc.paragraphs)
        
        for i in range(0, total_paragraphs, chunk_size):
            chunk_end = min(i + chunk_size, total_paragraphs)
            
            for j in range(i, chunk_end):
                if j < len(doc.paragraphs):
                    paragraph = doc.paragraphs[j]
                    if text in paragraph.text:
                        paragraph.text = paragraph.text.replace(text, '')
                        # Add strikethrough formatting for deleted text
                        for run in paragraph.runs:
                            run.font.strike = True
    
    def _add_clause(self, doc: DocxDocument, clause_text: str):
        """Add a new clause to the document"""
        new_paragraph = doc.add_paragraph(clause_text)
        for run in new_paragraph.runs:
            # Professional redlining for new clauses
            run.font.underline = True
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red for additions
    
    def _add_clause_chunked(self, doc: DocxDocument, clause_text: str):
        """Add a new clause to large documents"""
        new_paragraph = doc.add_paragraph(clause_text)
        for run in new_paragraph.runs:
            # Professional redlining for new clauses
            run.font.underline = True
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red for additions
    
    def _apply_firm_details(self, doc: DocxDocument, firm_details: Dict[str, Any]):
        """Apply firm details to signature blocks and firm information"""
        # Find and replace placeholder text with firm details
        for paragraph in doc.paragraphs:
            if '[FIRM_NAME]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[FIRM_NAME]', firm_details.get('name', ''))
                # Add professional redlining to show the change
                for run in paragraph.runs:
                    run.font.underline = True
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red for additions
            if '[FIRM_ADDRESS]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[FIRM_ADDRESS]', firm_details.get('address', ''))
                for run in paragraph.runs:
                    run.font.underline = True
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red for additions
            if '[SIGNER_NAME]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[SIGNER_NAME]', firm_details.get('signerName', ''))
                for run in paragraph.runs:
                    run.font.underline = True
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red for additions
            if '[SIGNER_TITLE]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[SIGNER_TITLE]', firm_details.get('signerTitle', ''))
                for run in paragraph.runs:
                    run.font.underline = True
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red for additions
    
    def _apply_firm_details_chunked(self, doc: DocxDocument, firm_details: Dict[str, Any]):
        """Apply firm details to large documents with chunked processing"""
        logger.info("Applying firm details to large document")
        chunk_size = 500
        total_paragraphs = len(doc.paragraphs)
        
        for i in range(0, total_paragraphs, chunk_size):
            chunk_end = min(i + chunk_size, total_paragraphs)
            
            for j in range(i, chunk_end):
                if j < len(doc.paragraphs):
                    paragraph = doc.paragraphs[j]
                    
                    if '[FIRM_NAME]' in paragraph.text:
                        paragraph.text = paragraph.text.replace('[FIRM_NAME]', firm_details.get('name', ''))
                        for run in paragraph.runs:
                            run.font.underline = True
                            run.font.color.rgb = RGBColor(255, 0, 0)
                    if '[FIRM_ADDRESS]' in paragraph.text:
                        paragraph.text = paragraph.text.replace('[FIRM_ADDRESS]', firm_details.get('address', ''))
                        for run in paragraph.runs:
                            run.font.underline = True
                            run.font.color.rgb = RGBColor(255, 0, 0)
                    if '[SIGNER_NAME]' in paragraph.text:
                        paragraph.text = paragraph.text.replace('[SIGNER_NAME]', firm_details.get('signerName', ''))
                        for run in paragraph.runs:
                            run.font.underline = True
                            run.font.color.rgb = RGBColor(255, 0, 0)
                    if '[SIGNER_TITLE]' in paragraph.text:
                        paragraph.text = paragraph.text.replace('[SIGNER_TITLE]', firm_details.get('signerTitle', ''))
                        for run in paragraph.runs:
                            run.font.underline = True
                            run.font.color.rgb = RGBColor(255, 0, 0)
            
            # Log progress for very large documents
            if total_paragraphs > 10000:
                progress = (chunk_end / total_paragraphs) * 100
                logger.info(f"Firm details progress: {progress:.1f}%")
    
    def _apply_signature(self, doc: DocxDocument, signature_path: str):
        """Apply signature image to the document"""
        try:
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            logger.info(f"Applying signature from: {signature_path}")
            
            # Find "Signed:" text and add signature right after it
            signature_added = False
            
            for paragraph in doc.paragraphs:
                if 'Signed:' in paragraph.text:
                    logger.info(f"Found 'Signed:' in paragraph: {paragraph.text}")
                    
                    # Clear the paragraph and rebuild it with proper signature placement
                    original_text = paragraph.text
                    paragraph.clear()
                    
                    # Split text around "Signed:"
                    parts = original_text.split('Signed:', 1)
                    
                    # Add text before "Signed:"
                    if parts[0].strip():
                        before_run = paragraph.add_run(parts[0])
                    
                    # Add "Signed:" text
                    signed_run = paragraph.add_run("Signed: ")
                    
                    # Add signature image immediately after "Signed:"
                    signature_run = paragraph.add_run()
                    signature_run.add_picture(signature_path, width=Inches(1.2))
                    
                    # Add any remaining text after "Signed:"
                    if len(parts) > 1 and parts[1].strip():
                        remaining_run = paragraph.add_run(parts[1])
                    
                    signature_added = True
                    logger.info("Signature added right next to 'Signed:' text")
                    break
            
            # If no "Signed:" found, look for signature placeholders
            if not signature_added:
                for paragraph in doc.paragraphs:
                    if '[SIGNATURE]' in paragraph.text:
                        # Replace placeholder with signature
                        paragraph.text = paragraph.text.replace('[SIGNATURE]', '')
                        
                        # Add signature image
                        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                        run.add_picture(signature_path, width=Inches(1.5))
                        
                        signature_added = True
                        logger.info("Signature added at placeholder location")
                        break
            
            # If still no signature added, add at the end
            if not signature_added:
                # Add a new paragraph for signature
                signature_paragraph = doc.add_paragraph()
                
                # Add "Signed:" text
                signed_run = signature_paragraph.add_run("Signed: ")
                
                # Add signature image right next to "Signed:"
                signature_run = signature_paragraph.add_run()
                signature_run.add_picture(signature_path, width=Inches(1.2))
                
                signature_added = True
                logger.info("Signature added at end of document")
            
            if signature_added:
                logger.info("Signature applied successfully")
            else:
                logger.warning("Could not find suitable location for signature")
            
        except Exception as e:
            logger.error(f"Error applying signature: {str(e)}")
            # Don't raise exception - signature is optional
    
    def apply_accepted_changes(self, document_path: str, accepted_changes: list, signature_path: str = None) -> dict:
        """
        Apply accepted changes to the document and optionally insert signature
        
        Args:
            document_path: Path to the original document
            accepted_changes: List of changes that were accepted
            signature_path: Optional path to signature image file
            
        Returns:
            dict with success status and output path
        """
        try:
            from docx import Document as DocxDocument
            from docx.shared import Inches
            import shutil
            
            # Load the original document
            doc = DocxDocument(document_path)
            logger.info(f"📄 Applying {len(accepted_changes)} accepted changes")
            
            # Only process changes if there are any
            if accepted_changes:
                # Build flat list of all paragraphs (including inside tables)
                all_paragraphs = list(doc.paragraphs)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            all_paragraphs.extend(cell.paragraphs)
                
                # Apply each accepted change
                for change in accepted_changes:
                    try:
                        current_text = change.get('current_text', '')
                        new_text = change.get('new_text', '')
                        
                        if not current_text or not new_text:
                            continue
                        applied = False
                        
                        # Replace text in paragraphs (with whitespace-aware matching)
                        for paragraph in all_paragraphs:
                            target_text = None
                            
                            if current_text in paragraph.text:
                                target_text = current_text
                            else:
                                target_text = self._find_text_with_whitespace(paragraph.text, current_text)
                            
                            if target_text:
                                if self._replace_text_in_paragraph(paragraph, target_text, new_text):
                                    logger.info(f"✅ Applied change: '{current_text[:30]}...' → '{new_text[:30]}...'")
                                    applied = True
                                    break
                        
                        if not applied:
                            logger.warning(f"⚠️ Could not apply change (text not found): '{current_text[:50]}'")
                                
                    except Exception as e:
                        logger.error(f"Error applying change: {e}")
                        continue
            else:
                logger.info("📄 No accepted changes to apply - returning document as-is (with signature if provided)")
            
            # Insert signature if provided
            if signature_path and os.path.exists(signature_path):
                logger.info(f"✍️ Inserting signature from: {signature_path}")
                self._insert_signature(doc, signature_path)
            
            # Generate output path
            output_path = self._generate_output_path(document_path)
            
            # Save the document
            doc.save(output_path)
            logger.info(f"💾 Saved redlined document to: {output_path}")
            
            # Clean up temporary signature file
            if signature_path and os.path.exists(signature_path):
                try:
                    os.remove(signature_path)
                    logger.info(f"🗑️ Cleaned up temporary signature file")
                except:
                    pass
            
            return {
                'success': True,
                'output_path': output_path,
                'message': f'Applied {len(accepted_changes)} changes successfully'
            }
            
        except Exception as e:
            logger.error(f"Error applying accepted changes: {str(e)}", exc_info=True)
            return {
                'success': False,
                'error': str(e)
            }
    
    def _insert_signature(self, doc, signature_path: str):
        """Insert signature image into the document signature block"""
        try:
            from docx.shared import Inches
            import re
            
            # Find the "Signed:" or "By:" field in the document
            signature_inserted = False
            
            for paragraph in doc.paragraphs:
                raw_text = paragraph.text
                text_lower = raw_text.strip().lower()
                
                # Look for signature line (typically after "Signed:" or near "By:")
                if text_lower.startswith('signed:') or (text_lower.startswith('by:') and '_' in raw_text):
                    # Find ALL underscore sequences so we can insert after the last one
                    underscore_matches = list(re.finditer(r'(_+)', raw_text))
                    
                    if underscore_matches:
                        last_match = underscore_matches[-1]
                        before_underscores = raw_text[:last_match.start()]
                        underscores = last_match.group(0)
                        after_underscores = raw_text[last_match.end():]
                        
                        # Clear paragraph and rebuild with precise ordering
                        paragraph.clear()
                        
                        if before_underscores:
                            paragraph.add_run(before_underscores)
                        
                        strike_run = paragraph.add_run(underscores)
                        strike_run.font.strike = True
                        
                        # Insert signature immediately after the underscores
                        sig_run = paragraph.add_run()
                        sig_run.add_picture(signature_path, width=Inches(2.0))
                        
                        if after_underscores:
                            paragraph.add_run(after_underscores)
                    else:
                        # No underscores found, just add at end with strikethrough on any existing underscores
                        for run in paragraph.runs:
                            if '_' in run.text:
                                run.font.strike = True
                        paragraph.add_run(' ')
                        sig_run = paragraph.add_run()
                        sig_run.add_picture(signature_path, width=Inches(2.0))
                    
                    signature_inserted = True
                    logger.info(f"✅ Inserted signature after underscores in: {raw_text[:40]}")
                    break
            
            if not signature_inserted:
                logger.warning("⚠️ Could not find signature location in document")
                
        except Exception as e:
            logger.error(f"Error inserting signature: {str(e)}", exc_info=True)
    
    def _generate_output_path(self, input_path: str) -> str:
        """Generate output path for the processed document"""
        # Create outputs directory if it doesn't exist
        output_dir = 'outputs'
        os.makedirs(output_dir, exist_ok=True)
        
        # Get the absolute path to the outputs directory
        base_dir = os.path.abspath(output_dir)
        filename = os.path.basename(input_path)
        name, ext = os.path.splitext(filename)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{name}_redlined_{timestamp}{ext}"
        return os.path.abspath(os.path.join(base_dir, output_filename))
